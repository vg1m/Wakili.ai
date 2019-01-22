import urllib.request
from bs4 import BeautifulSoup

#Milimani Civil Division
CauseList_page = 'http://www.kenyalaw.org/kl/index.php?id=8281'
page = urllib.request.urlopen(CauseList_page)
soup = BeautifulSoup(page, "lxml")
content=soup.find("div", attrs={"class": "page-content"})

import docx
doc = docx.Document()
doc.add_paragraph(content.text)
doc.save('Milimani_Civil_Causelist.docx')

#Milimani Criminal Division
CauseList_page = 'http://www.kenyalaw.org/kl/index.php?id=8282'
page = urllib.request.urlopen(CauseList_page)
soup = BeautifulSoup(page, "lxml")
content=soup.find("div", attrs={"class": "page-content"})

import docx
doc = docx.Document()
doc.add_paragraph(content.text)
doc.save('Milimani_Criminal_Causelist.docx')

#Milimani Family Division
CauseList_page = 'http://www.kenyalaw.org/kl/index.php?id=8283'
page = urllib.request.urlopen(CauseList_page)
soup = BeautifulSoup(page, "lxml")
content=soup.find("div", attrs={"class": "page-content"})

import docx
doc = docx.Document()
doc.add_paragraph(content.text)
doc.save('Milimani_Family_Causelist.docx')

#Milimani Commercial and Tax Division
CauseList_page = 'http://www.kenyalaw.org/kl/index.php?id=1592'
page = urllib.request.urlopen(CauseList_page)
soup = BeautifulSoup(page, "lxml")
content=soup.find("div", attrs={"class": "page-content"})

import docx
doc = docx.Document()
doc.add_paragraph(content.text)
doc.save('Milimani_Commercial&Tax_Causelist.docx')

#Milimani Anti-Corruption Division
CauseList_page = 'http://www.kenyalaw.org/kl/index.php?id=8996'
page = urllib.request.urlopen(CauseList_page)
soup = BeautifulSoup(page, "lxml")
content=soup.find("div", attrs={"class": "page-content"})

import docx
doc = docx.Document()
doc.add_paragraph(content.text)
doc.save('Milimani_ACEC_Causelist.docx')
