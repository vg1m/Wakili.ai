#Understand the target audience;Understand the natural language in which communication happens; 
#Understand the intent of the user;
#Come up with responses that can answer the user and give further clues

from bs4 import BeautifulSoup
import pdb
import requests
from parse import *
import re
import calendar
import threading
import random
import sys
import platform, os
import idna
import textwrap3
import socket
import urllib.request
from urllib.request import Request, urlopen
from urllib.error import URLError, HTTPError
from win10toast import ToastNotifier
import newspaper
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
import time
from time import sleep
from datetime import datetime, date


now = datetime.now()
toaster = ToastNotifier()

print("LegiBot [Version 1.0.2] | © 2019 All rights reserved |", date.today())
toaster.show_toast("Launching LegiBot",
                            "Your AI Lawyer. ",
                            icon_path="scrapercon.ico",
                            duration=5)
print("\n")
req = Request("http://kenyalaw.org/caselaw/")
try:
    response = urlopen(req)
except URLError as e:
    toaster.show_toast("No connection established :( ",
                                "Either your internet connection is off or KLR is down. Close and try again.",
                                icon_path="scrapercon.ico",
                                duration=5)
    time.sleep(6)
    sys.exit()
    time.sleep(0.5)

#Greet the user
User_Name=input("Name: ")
print("\n")
if now.hour == 1 or now.hour == 2 or now.hour == 3 or now.hour == 4 or now.hour == 5 or now.hour == 6 or now.hour == 7 or now.hour == 8 or now.hour == 9:
    print("Good morning", User_Name, "...")
elif now.hour==10 or now.hour == 11:
    print("Good morning", User_Name, "...")
elif now.hour==13 or now.hour == 12 or now.hour == 14 or now.hour == 15 or now.hour == 16:
    print("Good afternoon", User_Name, "...")
elif now.hour==17 or now.hour == 18 or now.hour == 19 or now.hour == 20 or now.hour == 21 or now.hour == 22 or now.hour == 23:
    print("Hey", User_Name, "going on overtime, eh?")
    

#Prompt user action
def main():
    print ("Pick an activity:- \n")
    print("1. Legislation.\n2. Legal Terms.\n3. Cases.\n4. Gazette Notices.\n5. Hansards.\n6. Cause lists ")
    print("\n")
    user=input("Activity: ")
    print("\n")
    if user=="6":
        court=input("1. High Court of Kenya.\nCourt: ")
        if court=="1":
            print("\nJurisdiction\n1. Milimani.\n ")
            region=input("Jurisdiction: ")
            if region =="1":
                print("\n")
                division=input("1. Civil.\n2. Criminal\n3. Family\n4. Commercial & Tax\n5. Anti-corruption and Economic Crimes\n\nDivision: ")
                if division=="1":
                    #Milimani Civil Division
                    CauseList_page = 'http://www.kenyalaw.org/kl/index.php?id=8281'
                    page = urllib.request.urlopen(CauseList_page)
                    soup = BeautifulSoup(page, "lxml")
                    content=soup.find("div", attrs={"class": "page-content"})

                    import docx
                    doc = docx.Document()
                    doc.add_paragraph(content.text)
                    doc.save('Milimani_Civil_Causelist.docx')
                    toaster.show_toast("Cause list saved. Check folder. ",
                                    " ",
                                    icon_path="scrapercon.ico",
                                    duration=5)
                    
                if division=="2":
                    #Milimani Criminal Division
                    CauseList_page = 'http://www.kenyalaw.org/kl/index.php?id=8282'
                    page = urllib.request.urlopen(CauseList_page)
                    soup = BeautifulSoup(page, "lxml")
                    content=soup.find("div", attrs={"class": "page-content"})

                    import docx
                    doc = docx.Document()
                    doc.add_paragraph(content.text)
                    doc.save('Milimani_Criminal_Causelist.docx')
                    toaster.show_toast("Cause list saved. Check folder. ",
                                    " ",
                                    icon_path="scrapercon.ico",
                                    duration=5)
                    
                if division=="3":
                    #Milimani Family Division
                    CauseList_page = 'http://www.kenyalaw.org/kl/index.php?id=8283'
                    page = urllib.request.urlopen(CauseList_page)
                    soup = BeautifulSoup(page, "lxml")
                    content=soup.find("div", attrs={"class": "page-content"})

                    import docx
                    doc = docx.Document()
                    doc.add_paragraph(content.text)
                    doc.save('Milimani_Family_Causelist.docx')
                    toaster.show_toast("Cause list saved. Check folder. ",
                                    " ",
                                    icon_path="scrapercon.ico",
                                    duration=5)
                    
                if division=="4":
                    #Milimani Commercial and Tax Division
                    CauseList_page = 'http://www.kenyalaw.org/kl/index.php?id=1592'
                    page = urllib.request.urlopen(CauseList_page)
                    soup = BeautifulSoup(page, "lxml")
                    content=soup.find("div", attrs={"class": "page-content"})

                    import docx
                    doc = docx.Document()
                    doc.add_paragraph(content.text)
                    doc.save('Milimani_Commercial&Tax_Causelist.docx')
                    toaster.show_toast("Cause list saved. Check folder. ",
                                    " ",
                                    icon_path="scrapercon.ico",
                                    duration=5)
                    
                if division=="5":
                    #Milimani Anti-Corruption Division
                    CauseList_page = 'http://www.kenyalaw.org/kl/index.php?id=8996'
                    page = urllib.request.urlopen(CauseList_page)
                    soup = BeautifulSoup(page, "lxml")
                    content=soup.find("div", attrs={"class": "page-content"})

                    import docx
                    doc = docx.Document()
                    doc.add_paragraph(content.text)
                    doc.save('Milimani_ACEC_Causelist.docx')
                    toaster.show_toast("Cause list saved. Check folder. ",
                                    " ",
                                    icon_path="scrapercon.ico",
                                    duration=5)

                    
                print("\n")
                print("-"*50)
                main()               
            
    if user=="5":
        year=input("Tell me something about the hansard publication, like the year and also a little context:   ")
        print("\n")
        print("Researching...")
        toaster.show_toast("Analysing your input against archive. ",
                                    "Please wait until results show. ",
                                    icon_path="scrapercon.ico",
                                    duration=5)
        chromedriver =  "./chromedriver.exe"
        options = Options()
        options.accept_untrusted_certs = True
        options.assume_untrusted_cert_issuer = True
        options.add_argument('disable-infobars')
        options.add_argument("--no-sandbox")
        options.add_argument("--allow-http-screen-capture")
        options.add_argument("--disable-impl-side-painting")
        options.add_argument("--disable-setuid-sandbox")
        options.add_argument("--disable-seccomp-filter-sandbox")
        options.add_argument("--disable-extensions")
        options.add_argument('--ignore-gpu-blacklist')
        options.add_argument('--no-default-browser-check')
        options.add_argument('--no-first-run')
        options.add_argument('--disable-default-apps')
        options.add_argument('--disable-infobars')
        options.add_argument('--disable-extensions')
        options.add_argument('--test-type')
        os.environ["webdriver.chrome.driver"] = chromedriver
        driver = webdriver.Chrome(chrome_options=options, executable_path= chromedriver)
        hansard_url = "https://books.google.co.ke/books/about/Kenya_National_Assembly_Official_Record.html?id=pvwVH2fQKWQC&hl=sw"
        driver.get(hansard_url)
        searchbox = driver.find_element_by_id("search_form_input")
        searchbox.send_keys(year)
        searchbox.send_keys("\n")
        print("\n")
        print("Continue researching from the browser. Once done, please close it.")
        print("\n")

        restart=input ("Pick a different search activity? [y/n] ").lower()
        if restart=="yes" or restart=="y":
            print("\n")
            main()
        else:
            print("\n")
            from datetime import datetime, date
            now = datetime.now()
            if date.today().weekday()== 0:
                if now.hour == 1 or now.hour == 2 or now.hour == 3 or now.hour == 4 or now.hour == 5 or now.hour == 6 or now.hour == 7 or now.hour == 8 or now.hour == 9:
                    print("Enjoy the rest of your Monday morning", User_Name, "------")
                if now.hour==10 or now.hour == 11 or now.hour == 12:
                    print("Have a good day", User_Name, )
                if now.hour==13 or now.hour == 14 or now.hour == 15 or now.hour == 16:
                    print("Have good afternoon", User_Name)
            if date.today().weekday()== 1:
                if now.hour == 1 or now.hour == 2 or now.hour == 3 or now.hour == 4 or now.hour == 5 or now.hour == 6 or now.hour == 7 or now.hour == 8 or now.hour == 9:
                    print("Enjoy the rest of your Monday morning", User_Name, "------")
                if now.hour==10 or now.hour == 11 or now.hour == 12:
                    print("Have a good day", User_Name, )
                if now.hour==13 or now.hour == 14 or now.hour == 15 or now.hour == 16:
                    print("Have good afternoon", User_Name)
            if date.today().weekday()== 2:
                if now.hour == 1 or now.hour == 2 or now.hour == 3 or now.hour == 4 or now.hour == 5 or now.hour == 6 or now.hour == 7 or now.hour == 8 or now.hour == 9:
                    print("Enjoy the rest of your Monday morning", User_Name, "------")
                if now.hour==10 or now.hour == 11 or now.hour == 12:
                    print("Have a good day", User_Name, )
                if now.hour==13 or now.hour == 14 or now.hour == 15 or now.hour == 16:
                    print("Have good afternoon", User_Name)
            if date.today().weekday()== 3:
                if now.hour == 1 or now.hour == 2 or now.hour == 3 or now.hour == 4 or now.hour == 5 or now.hour == 6 or now.hour == 7 or now.hour == 8 or now.hour == 9:
                    print("Enjoy the rest of your Monday morning", User_Name, "------")
                if now.hour==10 or now.hour == 11 or now.hour == 12:
                    print("Have a good day", User_Name, )
                if now.hour==13 or now.hour == 14 or now.hour == 15 or now.hour == 16:
                    print("Have good afternoon", User_Name)
            if date.today().weekday()== 4:
                if now.hour == 1 or now.hour == 2 or now.hour == 3 or now.hour == 4 or now.hour == 5 or now.hour == 6 or now.hour == 7 or now.hour == 8 or now.hour == 9:
                    print("Enjoy the rest of your morning", User_Name, "------")
                if now.hour==10 or now.hour == 11 or now.hour == 12:
                    print("Have a good day", User_Name, )
                if now.hour==13 or now.hour == 14 or now.hour == 15 or now.hour == 16:
                    print("Have good afternoon", User_Name)
                if now.hour==17 or now.hour == 18 or now.hour == 19 or now.hour == 20 or now.hour == 21 or now.hour == 22 or now.hour == 23:
                    print("Have a great weekend", User_Name)
            if date.today().weekday()== 5:
                if now.hour == 1 or now.hour == 2 or now.hour == 3 or now.hour == 4 or now.hour == 5 or now.hour == 6 or now.hour == 7 or now.hour == 8 or now.hour == 9:
                    print("Enjoy the rest of your morning", User_Name, "------")
                if now.hour==10 or now.hour == 11 or now.hour == 12:
                    print("Have a good day", User_Name, )
                if now.hour==13 or now.hour == 14 or now.hour == 15 or now.hour == 16:
                    print("Have good afternoon", User_Name)
                if now.hour==17 or now.hour == 18 or now.hour == 19 or now.hour == 20 or now.hour == 21 or now.hour == 22 or now.hour == 23:
                    print("Enjoy your weekend", User_Name)
            if date.today().weekday()== 6:
                if now.hour == 1 or now.hour == 2 or now.hour == 3 or now.hour == 4 or now.hour == 5 or now.hour == 6 or now.hour == 7 or now.hour == 8 or now.hour == 9:
                    print("Enjoy the rest of your morning", User_Name, "------")
                if now.hour==10 or now.hour == 11 or now.hour == 12:
                    print("Have a good day", User_Name, )
                if now.hour==13 or now.hour == 14 or now.hour == 15 or now.hour == 16:
                    print("Have good afternoon", User_Name)
                if now.hour==17 or now.hour == 18 or now.hour == 19 or now.hour == 20 or now.hour == 21 or now.hour == 22 or now.hour == 23:
                    print("Have a great week ahead", User_Name)
        print("Continue researching from the browser. Once done, please close it.")
        print("\n")
        print ("Activity Summary: ")
        print("> Researcher: ", User_Name)
        print("> Hansard Year: ", year)
        print("> Context: ", context)
        
        import csv
        #open a csv file with append, so old data will not be erased
        with open("Case_Search_Activity.csv", "a", encoding="utf-8") as csv_file:
            writer = csv.writer(csv_file)
            writer.writerow(["Date", "Researcher", "Hansard_Search", "Context"])
            Date=date.today()
            Researcher=User_Name
            Hansard_Search=year
            writer.writerow([Date, Researcher, Hansard_Search])           
        print("\n")
        main()
     
    if user=="4":
        #If day of the week is Friday, and time is 2 pm, download the Kenya Gazette.
    # If you want hours and minutes, add > if now.hour == 14 and now.minute == 12
        from datetime import datetime, date
        now = datetime.now()
        if date.today().weekday()== 4:
            wkg = ["http://kenyalaw.org/kenya_gazette/"]
            for urls in wkg:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
                for link in soup.findAll('a', attrs={'href': re.compile("^http://kenyalaw.org/kenya_gazette/gazette/download/")}):
                    Kenya_Gazette = link.get('href')
                    r = requests.get(Kenya_Gazette, allow_redirects=True)
                    open('Kenya Gazette.pdf', 'wb').write(r.content)
                    toaster.show_toast("Kenya Gazette Download complete ",
                                                "Check folder... ",
                                                icon_path="scrapercon.ico",
                                                duration=5)
            print(User_Name, "'s", "LegiBot Activity Summary: ")
            print("> Researcher: ", User_Name)
            print("> Query: Kenya Gazette ")
            print("-"*50)
            print("\n")
            if __name__ == "__main__":
                main()
        else:
            print("\n")
            print("For the latest issue please check this coming Friday as from 12 PM. For past issues, go to http://kenyalaw.org/kenya_gazette/. ")
            print("\n")
            if __name__ == "__main__":
                main()
                
    
    if user=="3":
        import time
        from datetime import datetime, date
        
        
        print("Starting case search..")
        print("\n")
        txt=input("1.) Which Law would you like cases on?  ")
        print("\n")
        Act=input("2.) Please include the section here. If you don't know, just skip: " )
        print("\n")
        FreeText=input("3.) Lastly, any free text? Perhaps a quote from a judge? ")
        print("\n")
        print("Researching...")
        print("\n")
        
        chromedriver =  "./chromedriver.exe"
        options = Options()
        options.accept_untrusted_certs = True
        options.assume_untrusted_cert_issuer = True
        options.add_argument('disable-infobars')
        options.add_argument("--no-sandbox")
        options.add_argument("--allow-http-screen-capture")
        options.add_argument("--disable-impl-side-painting")
        options.add_argument("--disable-setuid-sandbox")
        options.add_argument("--disable-seccomp-filter-sandbox")
        options.add_argument("--disable-extensions")
        options.add_argument('--ignore-gpu-blacklist')
        options.add_argument('--no-default-browser-check')
        options.add_argument('--no-first-run')
        options.add_argument('--disable-default-apps')
        options.add_argument('--disable-infobars')
        options.add_argument('--disable-extensions')
        options.add_argument('--test-type')
        os.environ["webdriver.chrome.driver"] = chromedriver
        driver = webdriver.Chrome(chrome_options=options, executable_path= chromedriver)
        cases_url = "http://kenyalaw.org/caselaw/"
        driver.get(cases_url)
        searchbox = driver.find_element_by_id("gsc-i-id1")
        searchbox.send_keys(txt,"," , Act,  ",", FreeText)
        searchbox.send_keys("\n")
        
        

        print("\n")
        restart=input ("Pick a different search activity? [y/n] ").lower()
        if restart=="yes" or restart=="y":
            print("\n")
            main()
        else:
            print("\n")
            import time
            from time import sleep
            from datetime import datetime, date
            now = datetime.now()
            if date.today().weekday()== 0:
                if now.hour == 1 or now.hour == 2 or now.hour == 3 or now.hour == 4 or now.hour == 5 or now.hour == 6 or now.hour == 7 or now.hour == 8 or now.hour == 9:
                    print("Enjoy the rest of your Monday morning", User_Name, "------")
                if now.hour==10 or now.hour == 11 or now.hour == 12:
                    print("Have a good day", User_Name, )
                if now.hour==13 or now.hour == 14 or now.hour == 15 or now.hour == 16:
                    print("Have good afternoon", User_Name)
            if date.today().weekday()== 1:
                if now.hour == 1 or now.hour == 2 or now.hour == 3 or now.hour == 4 or now.hour == 5 or now.hour == 6 or now.hour == 7 or now.hour == 8 or now.hour == 9:
                    print("Enjoy the rest of your Monday morning", User_Name, "------")
                if now.hour==10 or now.hour == 11 or now.hour == 12:
                    print("Have a good day", User_Name, )
                if now.hour==13 or now.hour == 14 or now.hour == 15 or now.hour == 16:
                    print("Have good afternoon", User_Name)
            if date.today().weekday()== 2:
                if now.hour == 1 or now.hour == 2 or now.hour == 3 or now.hour == 4 or now.hour == 5 or now.hour == 6 or now.hour == 7 or now.hour == 8 or now.hour == 9:
                    print("Enjoy the rest of your Monday morning", User_Name, "------")
                if now.hour==10 or now.hour == 11 or now.hour == 12:
                    print("Have a good day", User_Name, )
                if now.hour==13 or now.hour == 14 or now.hour == 15 or now.hour == 16:
                    print("Have good afternoon", User_Name)
            if date.today().weekday()== 3:
                if now.hour == 1 or now.hour == 2 or now.hour == 3 or now.hour == 4 or now.hour == 5 or now.hour == 6 or now.hour == 7 or now.hour == 8 or now.hour == 9:
                    print("Enjoy the rest of your Monday morning", User_Name, "------")
                if now.hour==10 or now.hour == 11 or now.hour == 12:
                    print("Have a good day", User_Name, )
                if now.hour==13 or now.hour == 14 or now.hour == 15 or now.hour == 16:
                    print("Have good afternoon", User_Name)
            if date.today().weekday()== 4:
                if now.hour == 1 or now.hour == 2 or now.hour == 3 or now.hour == 4 or now.hour == 5 or now.hour == 6 or now.hour == 7 or now.hour == 8 or now.hour == 9:
                    print("Enjoy the rest of your morning", User_Name, "------")
                if now.hour==10 or now.hour == 11 or now.hour == 12:
                    print("Have a good day", User_Name, )
                if now.hour==13 or now.hour == 14 or now.hour == 15 or now.hour == 16:
                    print("Have good afternoon", User_Name)
                if now.hour==17 or now.hour == 18 or now.hour == 19 or now.hour == 20 or now.hour == 21 or now.hour == 22 or now.hour == 23:
                    print("Have a great weekend", User_Name)
            if date.today().weekday()== 5:
                if now.hour == 1 or now.hour == 2 or now.hour == 3 or now.hour == 4 or now.hour == 5 or now.hour == 6 or now.hour == 7 or now.hour == 8 or now.hour == 9:
                    print("Enjoy the rest of your morning", User_Name, "------")
                if now.hour==10 or now.hour == 11 or now.hour == 12:
                    print("Have a good day", User_Name, )
                if now.hour==13 or now.hour == 14 or now.hour == 15 or now.hour == 16:
                    print("Have good afternoon", User_Name)
                if now.hour==17 or now.hour == 18 or now.hour == 19 or now.hour == 20 or now.hour == 21 or now.hour == 22 or now.hour == 23:
                    print("Enjoy your weekend", User_Name)
            if date.today().weekday()== 6:
                if now.hour == 1 or now.hour == 2 or now.hour == 3 or now.hour == 4 or now.hour == 5 or now.hour == 6 or now.hour == 7 or now.hour == 8 or now.hour == 9:
                    print("Enjoy the rest of your morning", User_Name, "------")
                if now.hour==10 or now.hour == 11 or now.hour == 12:
                    print("Have a good day", User_Name, )
                if now.hour==13 or now.hour == 14 or now.hour == 15 or now.hour == 16:
                    print("Have good afternoon", User_Name)
                if now.hour==17 or now.hour == 18 or now.hour == 19 or now.hour == 20 or now.hour == 21 or now.hour == 22 or now.hour == 23:
                    print("Have a great week ahead", User_Name)
                    
            
        print("Continue researching from the browser. Once done, please close it.")
        print("\n")
        print ("Activity Summary: ")
        print("> Researcher: ", User_Name)
        print("> Reference law: ", "'", txt, "'")
        print("> Specific section: ", "'",Act, "'")
        print("> Context: ", "'", FreeText, "'")
        
        import csv
        #open a csv file with append, so old data will not be erased
        with open("Case_Search_Activity.csv", "a", encoding="utf-8") as csv_file:
            writer = csv.writer(csv_file)
            writer.writerow(["Date", "Researcher", "Law", "Section", "Context"])
            Date=date.today()
            Researcher=User_Name
            Law=txt
            Section=Act
            Context=FreeText
            writer.writerow([Date, Researcher, Law, Section, Context])           
        print("\n")   

    if user=="2":
        import time
        from datetime import datetime, date
        import csv
        LegalTerm=input("Enter legal term: ").lower()
        with open('blacksLaw.csv', mode='r') as f:
            reader = csv.reader(f)
            for num, row in enumerate(reader):
                if LegalTerm in row[1]:
                    print ("\n", num, row)
        f.close()
        
        print("-"*50)
        restart=input ("Pick a different search activity? [y/n] ").lower()
        if restart=="yes" or restart=="y":
            print("\n")
            main()
        else:
            print("\n")
            from datetime import datetime, date
            now = datetime.now()
            if date.today().weekday()== 0:
                if now.hour == 1 or now.hour == 2 or now.hour == 3 or now.hour == 4 or now.hour == 5 or now.hour == 6 or now.hour == 7 or now.hour == 8 or now.hour == 9:
                    print("Enjoy the rest of your Monday morning", User_Name, "------")
                if now.hour==10 or now.hour == 11 or now.hour == 12:
                    print("Have a good day", User_Name, )
                if now.hour==13 or now.hour == 14 or now.hour == 15 or now.hour == 16:
                    print("Have good afternoon", User_Name)
            if date.today().weekday()== 1:
                if now.hour == 1 or now.hour == 2 or now.hour == 3 or now.hour == 4 or now.hour == 5 or now.hour == 6 or now.hour == 7 or now.hour == 8 or now.hour == 9:
                    print("Enjoy the rest of your Monday morning", User_Name, "------")
                if now.hour==10 or now.hour == 11 or now.hour == 12:
                    print("Have a good day", User_Name, )
                if now.hour==13 or now.hour == 14 or now.hour == 15 or now.hour == 16:
                    print("Have good afternoon", User_Name)
            if date.today().weekday()== 2:
                if now.hour == 1 or now.hour == 2 or now.hour == 3 or now.hour == 4 or now.hour == 5 or now.hour == 6 or now.hour == 7 or now.hour == 8 or now.hour == 9:
                    print("Enjoy the rest of your Monday morning", User_Name, "------")
                if now.hour==10 or now.hour == 11 or now.hour == 12:
                    print("Have a good day", User_Name, )
                if now.hour==13 or now.hour == 14 or now.hour == 15 or now.hour == 16:
                    print("Have good afternoon", User_Name)
            if date.today().weekday()== 3:
                if now.hour == 1 or now.hour == 2 or now.hour == 3 or now.hour == 4 or now.hour == 5 or now.hour == 6 or now.hour == 7 or now.hour == 8 or now.hour == 9:
                    print("Enjoy the rest of your Monday morning", User_Name, "------")
                if now.hour==10 or now.hour == 11 or now.hour == 12:
                    print("Have a good day", User_Name, )
                if now.hour==13 or now.hour == 14 or now.hour == 15 or now.hour == 16:
                    print("Have good afternoon", User_Name)
            if date.today().weekday()== 4:
                if now.hour == 1 or now.hour == 2 or now.hour == 3 or now.hour == 4 or now.hour == 5 or now.hour == 6 or now.hour == 7 or now.hour == 8 or now.hour == 9:
                    print("Enjoy the rest of your morning", User_Name, "------")
                if now.hour==10 or now.hour == 11 or now.hour == 12:
                    print("Have a good day", User_Name, )
                if now.hour==13 or now.hour == 14 or now.hour == 15 or now.hour == 16:
                    print("Have good afternoon", User_Name)
                if now.hour==17 or now.hour == 18 or now.hour == 19 or now.hour == 20 or now.hour == 21 or now.hour == 22 or now.hour == 23:
                    print("Have a great weekend", User_Name)
            if date.today().weekday()== 5:
                if now.hour == 1 or now.hour == 2 or now.hour == 3 or now.hour == 4 or now.hour == 5 or now.hour == 6 or now.hour == 7 or now.hour == 8 or now.hour == 9:
                    print("Enjoy the rest of your morning", User_Name, "------")
                if now.hour==10 or now.hour == 11 or now.hour == 12:
                    print("Have a good day", User_Name, )
                if now.hour==13 or now.hour == 14 or now.hour == 15 or now.hour == 16:
                    print("Have good afternoon", User_Name)
                if now.hour==17 or now.hour == 18 or now.hour == 19 or now.hour == 20 or now.hour == 21 or now.hour == 22 or now.hour == 23:
                    print("Enjoy your weekend", User_Name)
            if date.today().weekday()== 6:
                if now.hour == 1 or now.hour == 2 or now.hour == 3 or now.hour == 4 or now.hour == 5 or now.hour == 6 or now.hour == 7 or now.hour == 8 or now.hour == 9:
                    print("Enjoy the rest of your morning", User_Name, "------")
                if now.hour==10 or now.hour == 11 or now.hour == 12:
                    print("Have a good day", User_Name, )
                if now.hour==13 or now.hour == 14 or now.hour == 15 or now.hour == 16:
                    print("Have good afternoon", User_Name)
                if now.hour==17 or now.hour == 18 or now.hour == 19 or now.hour == 20 or now.hour == 21 or now.hour == 22 or now.hour == 23:
                    print("Have a great week ahead", User_Name)
            
            print("\n")
            print ("Activity Summary: ")
            print("> Researcher: ", User_Name)
            print("> Black's Law Context Search: ", "'", LegalTerm, "'")
            import csv
            #open a csv file with append, so old data will not be erased
            with open("Blacks_Law_Context_Search.csv", "a", encoding="utf-8") as csv_file:
                writer = csv.writer(csv_file)
                writer.writerow(["Date", "Researcher", "Query"])
                Date=date.today()
                Researcher=User_Name
                Query=LegalTerm
                writer.writerow([Date, Researcher, Query])
                print("\n")
            input("Press 'Enter' to exit ")
            sys.exit()
        print("\n")
        if __name__ == "__main__":
            main()
    

    if user=="1":
        import time
        
        actTitle= input("Free Text: ").lower()
        toaster.show_toast("Scraping in progress...",
                                    " ",
                                    icon_path="scrapercon.ico",
                                    duration=5)

        if actTitle=="Advocates" or actTitle=="advocates" or actTitle=="lawyers" or actTitle=="lawyer" or actTitle=="advocate" or actTitle=="counsel" or actTitle=="esquire":
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%2016"]
            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")

            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            quote_page = "http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%2016"
            page = urllib.request.urlopen(quote_page)
            soup = BeautifulSoup(page, "lxml")
            Amending_Acts=soup.find("div", attrs={"id": "longTitle_1_V16_t16"})
            print(Amending_Acts.text)
                
            time = soup.find("select", id="cboPIT")
            print("Amendment History")
            print(time.text)
            print("\n")    
                        
            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/A/Advocates%20Act%20Cap.%2016%20-%20No.%2018%20of%201989/docs/AdvocatesAct18of1989.pdf"
            r = requests.get(Download, allow_redirects=True)
            toaster.show_toast("Download Complete ",
                                    "",
                                    icon_path="scrapercon.ico",
                                    duration=5)     
            open('Advocates_Act .pdf', 'wb').write(r.content)
            print("-"*50)
            advo3=input ("Get subsidiary legislation? [y/n] ").lower()
            if advo3=="yes" or advo3=="y":
                advo3= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/A/Advocates%20Act%20Cap.%2016%20-%20No.%2018%20of%201989/subsidiary%20legislation/docs/AdvocatesAct18of1989_subsidiary.pdf"
                r = requests.get(advo3, allow_redirects=True)
                toaster.show_toast("Download Complete ",
                                    "Check folder for *Advocates_Regulations* ",
                                    icon_path="scrapercon.ico",
                                    duration=5)   
            open('Advocates_Regulations.pdf', 'wb').write(r.content)
            print("-"*50)
            print("\n")
            print("-"*50)
            hansard=input("Get Parliamentary Debates? [y/n] " ).lower()
            if hansard=="y":
                hansard1= "http://www.parliament.go.ke/the-national-assembly/house-business/hansard/item/download/3501_214de0bec5e6d1836333bd9eeac25c66"
                r = requests.get(hansard1, allow_redirects=True)
                open('Advocates Act Hansard.pdf', 'wb').write(r.content)
                toaster.show_toast("Download Complete ",
                                            "Check folder for *Advocates_Act_Hansard* ",
                                            icon_path="scrapercon.ico",
                                            duration=5)        
    ###
        if actTitle=="Banking" or actTitle=="bank" or actTitle=="banking":
            
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%20488"]
            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
       
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)
     
            quote_page = "http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%20488"
            page = urllib.request.urlopen(quote_page)
            soup = BeautifulSoup(page, "lxml")
            Amending_Acts=soup.find("div", attrs={"id": "longTitle_1_V16_t16"})
            print(Amending_Acts.text)
                
            time = soup.find("select", id="cboPIT")
            print("Amendment History")
            print(time.text)
            print("\n")

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/B/Banking%20Act%20Cap.%20488%20-%20No.%209%20of%201989/docs/BankingAct9of1989.pdf"
            r = requests.get(Download, allow_redirects=True)
            toaster.show_toast("Download Complete ",
                                    "",
                                    icon_path="scrapercon.ico",
                                    duration=5)
            open('Banking .pdf', 'wb').write(r.content)
            print("-"*50)
            hansard=input("Fetch Parliamentary Debates? [y/n] ").lower()
            if hansard=="y":
                hansard1= "http://www.parliament.go.ke/the-national-assembly/house-business/hansard/item/download/2279_9aa3eaf814794aa12b4df6680fb9dee8"
                r = requests.get(hansard1, allow_redirects=True)
                open('Banking Act Hansard.pdf', 'wb').write(r.content)
                toaster.show_toast("Download Complete ",
                                        "Saved as *Banking Act Hansard",
                                        icon_path="scrapercon.ico",
                                        duration=5)
        
    ###        
        if actTitle=="Accountants" or actTitle=="accountants" or actTitle=="accountant" or actTitle=="accounts":
         
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=No.%2015%20of%202008"]
            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
           
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            quote_page = "http://www.kenyalaw.org/lex//actview.xql?actid=No.%2015%20of%202008"
            page = urllib.request.urlopen(quote_page)
            soup = BeautifulSoup(page, "lxml")
            Amending_Acts=soup.find("div", attrs={"id": "longTitle_1_V4_t4"})
            print(Amending_Acts.text)
            
            time = soup.find("select", id="cboPIT")
            print("Amendment History")
            print(time.text)
            
            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/A/Accountants%20Act%20Cap.%20531%20-%20No.%2015%20of%202008/docs/AccountantsAct15of2008.pdf"
            r = requests.get(Download, allow_redirects=True)
            toaster.show_toast("Download Complete ",
                                    "",
                                    icon_path="scrapercon.ico",
                                    duration=5)
            open('Accountants .pdf', 'wb').write(r.content)
            print("-"*50)
            hansard=input("Fetch Parliamentary Debates? [y/n] ").lower()
            if hansard=="y":
                hansard1= "https://books.google.co.ke/books?id=Mv0MboYdOU0C&pg=PA338&lpg=PA338&dq=kenya+accountants+act+%22hansard%22&source=bl&ots=zFMKwDuvk9&sig=PbHcaSVFNN-KS4R52MsN3Tp3r7Y&hl=en&sa=X&ved=2ahUKEwiwhJ6708bcAhUOhRoKHfNxBToQ6AEwCHoECAgQAQ#v=onepage&q=kenya%20accountants%20act%20%22hansard%22&f=false"
                r = requests.get(hansard1, allow_redirects=True)
                open('Accountants Act Hansard.html', 'wb').write(r.content)
                toaster.show_toast("Download Complete ",
                                        "Saved as *Accountants Act Hansard* ",
                                        icon_path="scrapercon.ico",
                                        duration=5)
    ###
        if actTitle=="African Development Bank" or actTitle=="African development bank" or actTitle=="african development bank" or actTitle=="ADB" or actTitle=="adb":
          
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%20492"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")

            
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            quote_page = "http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%20492"
            page = urllib.request.urlopen(quote_page)
            soup = BeautifulSoup(page, "lxml")
            Amending_Acts=soup.find("div", attrs={"class": "act-long-title"})
            print(Amending_Acts.text)
                
            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/A/African%20Development%20Bank%20Act%20Cap.%20492%20-%20No%204%20of%201964/docs/AfricanDevelopmentBankAct4of1964.pdf"
            r = requests.get(Download, allow_redirects=True)
            toaster.show_toast("Downloaded Online Version ",
                                        " ",
                                        icon_path="scrapercon.ico",
                                        duration=5)
            open('African Devlopment Bank .pdf', 'wb').write(r.content)
    ###
        if actTitle=="Age of majority" or actTitle=="age of majority" or actTitle=="Age Of Majority" or actTitle=="Age of Majority" or actTitle=="majority age" or actTitle=="Majority Age" or actTitle=="Majority age":
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%2033" ]
            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
            
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)
        
            quote_page = "http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%2033"
            page = urllib.request.urlopen(quote_page)
            soup = BeautifulSoup(page, "lxml")
            Amending_Acts=soup.find("div", attrs={"class": "act-long-title"})
            print(Amending_Acts.text)
            print("\n")
            
            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/A/Age%20of%20Majority%20Act%20Cap.%2033%20-%20Act%20No.%201%20of%201974/docs/AgeofMajorityActAct1of1974.pdf"
            r = requests.get(Download, allow_redirects=True)
            toaster.show_toast("Downloaded Online Version ",
                                        " ",
                                        icon_path="scrapercon.ico",
                                        duration=5)
            open('Age of majority .pdf', 'wb').write(r.content)
            print("-"*50)
            hansard=input("Fetch Parliamentary Debates on? [y/n] ").lower()
            if hansard=="y":
                hansard1= "https://books.google.co.ke/books?id=-opPV74o1aUC&pg=PA252&lpg=PA252&dq=kenya+Age+of+majority+Act+hansard&source=bl&ots=sm-UM6BDwi&sig=TZoDPs_BOnnjw_9IgM6ad38K3bc&hl=en&sa=X&ved=2ahUKEwj8kdbv2cbcAhWhxIUKHVfNCrM4ChDoATACegQIABAB#v=onepage&q=kenya%20Age%20of%20majority%20Act%20hansard&f=false"
                r = requests.get(hansard1, allow_redirects=True)
                hansard2= "https://books.google.co.ke/books?id=8ubVSfGcxNoC&pg=PA379&lpg=PA379&dq=kenya+Age+of+majority+Act+hansard&source=bl&ots=dWfOH2Rlbc&sig=GB9Z7Q5E6dEHN-L6EWYFfg7sbLw&hl=en&sa=X&ved=2ahUKEwj8kdbv2cbcAhWhxIUKHVfNCrM4ChDoATABegQIAhAB#v=onepage&q=kenya%20Age%20of%20majority%20Act%20hansard&f=false"
                r = requests.get(hansard2, allow_redirects=True)
                open('Age of Majority Act Hansard1933.html', 'wb').write(r.content)
                open('Age of Majority Act Hansard1974.html', 'wb').write(r.content)
                print("-----Downloads complete-----")
                toaster.show_toast("Download Complete ",
                                        "Saved as *Age of Majority Act Hansard* ",
                                        icon_path="scrapercon.ico",
                                        duration=5)
    ###        
        if actTitle=="Agricultural development corporation" or actTitle=="ADC" or actTitle=="adc":
            
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%20444"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
       
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            quote_page = "http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%20444"
            page = urllib.request.urlopen(quote_page)
            soup = BeautifulSoup(page, "lxml")
            Amending_Acts=soup.find("div", attrs={"class": "act-long-title"})
            print(Amending_Acts.text)
            
            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/A/Agricultural%20Development%20Corporation%20Act%20Cap.%20444%20-%20No.%207%20of%201965/docs/AgriculturalDevelopmentCorporationAct7of1965.pdf"
            r = requests.get(Download, allow_redirects=True)
            toaster.show_toast("Downloaded Online Version ",
                                        " ",
                                        icon_path="scrapercon.ico",
                                        duration=5)
            open('ADC.pdf', 'wb').write(r.content)
            print("-"*50)
            hansard=input("Fetch Parliamentary Debates? [y/n] ").lower()
            if hansard=="y":
                hansard1= "https://books.google.co.ke/books?id=r4DX0_UKLoYC&pg=PA1434&lpg=PA1434&dq=Agricultural+development+corporation++Act+kenya+assembly+hansard&source=bl&ots=aoIHVM7fJy&sig=VxTZrmY90CxFswXAXlRZ8m2qL94&hl=en&sa=X&ved=2ahUKEwjesujN3sbcAhUDRBoKHbaEBx4Q6AEwCXoECAkQAQ#v=onepage&q=Agricultural%20development%20corporation%20%20Act%20kenya%20assembly%20hansard&f=false"
                r = requests.get(hansard1, allow_redirects=True)
                hansard2= "https://books.google.co.ke/books?id=7ozb_Mmxy7MC&pg=RA1-PA1564&lpg=RA1-PA1564&dq=Agricultural+development+corporation++Act+kenya+assembly+hansard&source=bl&ots=FxH42kVsn8&sig=vQA-D0tfQfGu8JvDwTV-0TMkmjo&hl=en&sa=X&ved=2ahUKEwjesujN3sbcAhUDRBoKHbaEBx4Q6AEwB3oECAcQAQ#v=onepage&q=Agricultural%20development%20corporation%20%20Act%20kenya%20assembly%20hansard&f=false"
                r = requests.get(hansard2, allow_redirects=True)
                open('Agricultural development corporation Act Hansard1981.html', 'wb').write(r.content)
                open('Agricultural development corporation Act Hansard1986.html', 'wb').write(r.content)
                print("-----Downloads complete-----")
                toaster.show_toast("Download Complete ",
                                        "Saved as *Agricultural Development Corporation Act Hansard* ",
                                        icon_path="scrapercon.ico",
                                        duration=5)
    ###
        if actTitle=="Insolvency" or actTitle=="insolvency":
            
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=No.%2018%20of%202015"]

            toaster.show_toast("Processing speed",
                                        "Some statutes are large in size and may be slow to process than others. ",
                                            icon_path="scrapercon.ico",
                                            duration=4)

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")

            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            quote_page = "http://www.kenyalaw.org/lex//actview.xql?actid=No.%2018%20of%202015"
            page = urllib.request.urlopen(quote_page)
            soup = BeautifulSoup(page, "lxml")
            Amending_Acts=soup.find("div", attrs={"id": "longTitle_1_V6_t6"})
            print(Amending_Acts.text)

            time = soup.find("select", id="cboPIT")
            print("Amendment history")
            print(time.text)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/I/Insolvency%20Act%20-%20No.%2018%20of%202015/docs/InsolvencyAct18of2015.pdf"
            r = requests.get(Download, allow_redirects=True)
            toaster.show_toast("Downloaded Online Version ",
                                        " ",
                                        icon_path="scrapercon.ico",
                                        duration=5)
            open('Insolvency .pdf', 'wb').write(r.content)
            
    ###
        if actTitle=="Agricultural and Food Authority" or actTitle=="AFA" or actTitle=="afa" or actTitle=="agricultural and food authority" :
           
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=No.%2013%20of%202013"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
         
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)
              
            quote_page = "http://www.kenyalaw.org/lex//actview.xql?actid=No.%2013%20of%202013"
            page = urllib.request.urlopen(quote_page)
            soup = BeautifulSoup(page, "lxml")
            Amending_Acts=soup.find("div", attrs={"id": "longTitle_1_V5_t5"})
            print(Amending_Acts.text)

            time = soup.find("select", id="cboPIT")
            print("Amendment history")
            print(time.text)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/A/Agriculture%20and%20Food%20Authority%20Act%20No.%2013%20of%202013/docs/Agriculture%20and%20Food%20Authority%20Act%20No.%2013%20of%202013.pdf"
            r = requests.get(Download, allow_redirects=True)
            toaster.show_toast("Downloaded Online Version ",
                                        " ",
                                        icon_path="scrapercon.ico",
                                        duration=5)
            
            open('Agricultural Food Authority .pdf', 'wb').write(r.content)
    ###
        if actTitle=="Air Passenger Service Charge" or actTitle=="air passenger service charge" or actTitle=="air passenger" or actTitle=="air" or actTitle=="air passenger service" :
           
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%20475"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
        
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            quote_page = "http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%20475"
            page = urllib.request.urlopen(quote_page)
            soup = BeautifulSoup(page, "lxml")
            Amending_Acts=soup.find("div", attrs={"id": "longTitle_1_V16_t16"})
            print(Amending_Acts.text)
                
            time = soup.find("select", id="cboPIT")
            print("Latest Edition")
            print(time.text)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/A/Air%20Passenger%20Service%20Charge%20Act%20Cap.%20475%20-%20No.%2021%20of%201970/docs/AirPassengerServiceChargeAct21of1970.pdf"
            r = requests.get(Download, allow_redirects=True)
            toaster.show_toast("Downloaded Online Version ",
                                        " ",
                                        icon_path="scrapercon.ico",
                                        duration=5)
            open('Air Passenger Service Charge .pdf', 'wb').write(r.content)
            print("-"*50)
            hansard=input("Fetch Parliamentary Debates? [y/n] ").lower()
            if hansard=="y":
                hansard1= "https://books.google.co.ke/books?id=g67aPKwJ9mAC&pg=PT253&lpg=PT253&dq=kenya+assembly+hansard+Air+Passenger+Service+Charge+Act&source=bl&ots=z0hJj0OMgL&sig=nwFiGZRrbPtl4q1HsOOocxKUpCU&hl=en&sa=X&ved=2ahUKEwiP-MrB48bcAhUSy4UKHZvICvwQ6AEwA3oECAMQAQ#v=onepage&q=kenya%20assembly%20hansard%20Air%20Passenger%20Service%20Charge%20Act&f=false"
                r = requests.get(hansard1, allow_redirects=True)
                hansard2= "https://books.google.co.ke/books?id=LwTcZFh6k6oC&pg=PT29&lpg=PT29&dq=kenya+assembly+hansard+Air+Passenger+Service+Charge+Act&source=bl&ots=Hx2aHLez0u&sig=h9xYi_eyp_jtJRXWs-m911iAAYA&hl=en&sa=X&ved=2ahUKEwiP-MrB48bcAhUSy4UKHZvICvwQ6AEwBXoECAUQAQ#v=onepage&q=kenya%20assembly%20hansard%20Air%20Passenger%20Service%20Charge%20Act&f=false"
                r = requests.get(hansard2, allow_redirects=True)
                hansard3= "https://books.google.co.ke/books?id=R10vOiTletUC&pg=PT28&lpg=PT28&dq=kenya+assembly+hansard+Air+Passenger+Service+Charge+Act&source=bl&ots=Ivy4MnJ8cA&sig=x7QdTMekiIDeBwoGTfwGkHwU08w&hl=en&sa=X&ved=2ahUKEwiP-MrB48bcAhUSy4UKHZvICvwQ6AEwB3oECAgQAQ#v=onepage&q=kenya%20assembly%20hansard%20Air%20Passenger%20Service%20Charge%20Act&f=false"
                r = requests.get(hansard3, allow_redirects=True)
                open('Air Passenger Service Charge Act Hansard1988.html', 'wb').write(r.content)
                open('Air Passenger Service Charge Act Hansard2004.html', 'wb').write(r.content)
                open('Air Passenger Service Charge Act Hansard1999.html', 'wb').write(r.content)
                toaster.show_toast("Download Complete ",
                                        "Saved as *Air Passenger Service Charge Act Hansards* ",
                                        icon_path="scrapercon.ico",
                                        duration=5)
    ###
        if actTitle=="Alcoholic Drinks Control" or actTitle=="alcoholic Control" or actTitle=="alcohol" or actTitle=="alcoholic drinks":
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=No.%204%20of%202010"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")

            
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            quote_page = "http://www.kenyalaw.org/lex//actview.xql?actid=No.%204%20of%202010"
            page = urllib.request.urlopen(quote_page)
            soup = BeautifulSoup(page, "lxml")
            Amending_Acts=soup.find("div", attrs={"class": "act-long-title"})
            print(Amending_Acts.text)

            time = soup.find("select", id="cboPIT")
            print("Amendment history")
            print(time.text)
           
            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/A/Alcoholic%20Drinks%20Control%20Act%20Cap.%20121A%20-%20No.%204%20of%202010/docs/AlcoholicDrinksControlAct4of2010.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            toaster.show_toast("Downloaded Online Version ",
                                        " ",
                                        icon_path="scrapercon.ico",
                                        duration=5)
            open('Alcoholic Drinks Control .pdf', 'wb').write(r.content)
            print("-"*50)
            hansard=input("Fetch Parliamentary Debates? [y/n] ").lower()
            if hansard=="y":
                hansard1= "https://books.google.co.ke/books?id=FeuFc7dvnAcC&pg=PT17&lpg=PT17&dq=kenya+assembly+hansard+Alcoholic+Drinks+Control++Act&source=bl&ots=pC1ZNVubsN&sig=YDBNG0Iqc9P6IaNT4m0AUZVqAGU&hl=en&sa=X&ved=2ahUKEwi4lbLl7cbcAhWlyIUKHYnbDzQ4ChDoATADegQIAxAB#v=onepage&q=kenya%20assembly%20hansard%20Alcoholic%20Drinks%20Control%20%20Act&f=false"
                r = requests.get(hansard1, allow_redirects=True)
                hansard2="https://books.google.co.ke/books?id=MUemcU8wPGMC&pg=PT23&lpg=PT23&dq=kenya+assembly+hansard+Alcoholic+Drinks+Control++Act&source=bl&ots=xi224SC0oz&sig=cwpVmqulcbCFF-a0hDvEQr2OTMM&hl=en&sa=X&ved=2ahUKEwi4lbLl7cbcAhWlyIUKHYnbDzQ4ChDoATAHegQIBxAB#v=onepage&q=kenya%20assembly%20hansard%20Alcoholic%20Drinks%20Control%20%20Act&f=false"
                r = requests.get(hansard2, allow_redirects=True)
                open('Alcoholic Drinks Control Hansard2010.html', 'wb').write(r.content)
                open('Alcoholic Drinks Control Hansard2010.html', 'wb').write(r.content)
                print("-----Downloads complete-----")
                toaster.show_toast("Download Complete ",
                                        "Saved as *Alcoholic Drinks Control Hansards* ",
                                        icon_path="scrapercon.ico",
                                        duration=5)
    ###
        if actTitle=="Anatomy" or actTitle=="anatomy":
          
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%20249"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")

            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            quote_page = "http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%20249"
            page = urllib.request.urlopen(quote_page)
            soup = BeautifulSoup(page, "lxml")
            Amending_Acts=soup.find("div", attrs={"id": "longTitle_1_V3_t3"})
            print(Amending_Acts)
            
##            time = soup.find("select", id="cboPIT")
##            print("Amendment history")
##            print(time.text)
            
            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/A/Anatomy%20Act%20Cap.%20249%20-%20No.%2021%20of%201967/docs/AnatomyAct21of1967.pdf"
            r = requests.get(Download, allow_redirects=True)
            toaster.show_toast("Downloaded Online Version ",
                                        " ",
                                        icon_path="scrapercon.ico",
                                        duration=5)
            open('Anatomy .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Animal Diseases" or actTitle=="animal diseases" or actTitle=="Animal diseases":
           
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%20364"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
            
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            quote_page = "http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%20364"
            page = urllib.request.urlopen(quote_page)
            soup = BeautifulSoup(page, "lxml")
            Amending_Acts=soup.find("div", attrs={"id": "longTitle_1_V3_t3"})
            print(Amending_Acts.text)
            
            time = soup.find("select", id="cboPIT")
            print("Amendment history")
            print(time.text)
            
            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/A/Animal%20Diseases%20Cap.%20364%20-%20Act%20No.%204%20of%201965/docs/AnimaldiseasesAct4of1965.pdf"
            r = requests.get(Download, allow_redirects=True)
            toaster.show_toast("Downloaded Online Version ",
                                        " ",
                                        icon_path="scrapercon.ico",
                                        duration=5)
            open('Animal Diseases .pdf', 'wb').write(r.content)
            print("-"*50)
            hansard=input("Fetch Parliamentary Debates? [y/n] ").lower()
            if hansard=="y":
                hansard1= "https://books.google.co.ke/books?id=gw4w99W4_7IC&pg=RA1-PA1468&lpg=RA1-PA1468&dq=Kenya+assembly+hansard+animal+diseases&source=bl&ots=RP7zJUUpkV&sig=KO-7Rgy-CxTGGtoUWYlLEZFaeeM&hl=en&sa=X&ved=2ahUKEwjE19z18s3cAhUEvxoKHZKhDRwQ6AEwB3oECAQQAQ#v=onepage&q=Kenya%20assembly%20hansard%20animal%20diseases&f=false"
                r = requests.get(hansard1, allow_redirects=True)
                hansard2= "https://books.google.co.ke/books?id=S9HweMSOY3MC&pg=PT9&lpg=PT9&dq=Kenya+assembly+hansard+animal+diseases&source=bl&ots=pLdNOAkmQR&sig=JJf1SDB-uqq8yFDruxg-OKltv_E&hl=en&sa=X&ved=2ahUKEwjE19z18s3cAhUEvxoKHZKhDRwQ6AEwBXoECAUQAQ#v=onepage&q=Kenya%20assembly%20hansard%20animal%20diseases&f=false"
                r = requests.get(hansard2, allow_redirects=True)
                open('Animal Diseases Hansard1995a.html', 'wb').write(r.content)
                open('Animal Diseases Hansard1995b.html', 'wb').write(r.content)
                toaster.show_toast("Downloads complete ",
                                        "Saved as *Animal Diseases Act Hansards ",
                                        icon_path="scrapercon.ico",
                                        duration=5) 

    ###
        if actTitle=="Animal Technicians" or actTitle=="animal technicians" or actTitle=="Animal technicians":
            
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=No.%2010%20of%202010"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")

            quote_page = "http://www.kenyalaw.org/lex//actview.xql?actid=No.%2010%20of%202010"
            page = urllib.request.urlopen(quote_page)
            soup = BeautifulSoup(page, "lxml")
            Amending_Acts=soup.find("div", attrs={"class": "act-long-title"})
            print(Amending_Acts.text)
            
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)
                
##            time = soup.find("select", id="cboPIT")
##            print("Amendment history")
##            print(time.text)  

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/A/Animal%20Technicians%20Act%20Cap.%20364A%20-%20No.%2011%20of%202010/docs/AnimalTechnicianAct11of2010.pdf"
            r = requests.get(Download, allow_redirects=True)
            toaster.show_toast("Downloaded Online Version ",
                                        " ",
                                        icon_path="scrapercon.ico",
                                        duration=5)
            open('Animal Technicians .pdf', 'wb').write(r.content)
            print("-"*50)
            hansard=input("Fetch Parliamentary Debates? [y/n] ").lower()
            if hansard=="y":
                hansard1= "http://info.mzalendo.com/hansard/sitting/national_assembly/2008-08-06-09-00-00"
                r = requests.get(hansard1, allow_redirects=True)
                hansard2= "https://books.google.co.ke/books?id=3yKdQH3-1hUC&pg=PT36&lpg=PT36&dq=Kenya+assembly+hansard+animal+technicians&source=bl&ots=OD30Wl9Ej1&sig=dEF47ekV9TTlnlifoZAVqkeFbJE&hl=en&sa=X&ved=2ahUKEwjd_cf08M3cAhUFQhoKHe6dBUIQ6AEwAXoECAkQAQ#v=onepage&q=Kenya%20assembly%20hansard%20animal%20technicians&f=false"
                r = requests.get(hansard2, allow_redirects=True)
                hansard3= "https://books.google.co.ke/books?id=mUV56KFDCF8C&pg=PT36&lpg=PT36&dq=Kenya+assembly+hansard+animal+technicians&source=bl&ots=kcR3zYTKNy&sig=E3VIruL_hjvwQW1UcFI2_d6Np_U&hl=en&sa=X&ved=2ahUKEwjd_cf08M3cAhUFQhoKHe6dBUIQ6AEwBHoECAYQAQ#v=onepage&q=Kenya%20assembly%20hansard%20animal%20technicians&f=false"
                r = requests.get(hansard3, allow_redirects=True)
                open('Animal Technicians Hansard2008.html', 'wb').write(r.content)
                open('Animal Technicians Hansard2008a.html', 'wb').write(r.content)
                open('Animal Technicians Hansard2010.html', 'wb').write(r.content)
                toaster.show_toast("Downloads complete ",
                                        "Saved as *Animal Technicians Act Hansards* ",
                                        icon_path="scrapercon.ico",
                                        duration=5)

    ###       
        if actTitle=="Anti-corruption and Economic Crimes" or actTitle=="ACEC" or actTitle=="acec" or actTitle=="anti-corruption " or actTitle=="anticorruption " or actTitle=="anti corruption " or actTitle=="corruption" or actTitle=="economic crimes ":
           
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=No.%203%20of%202003"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
          
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)
                    
            quote_page = "http://www.kenyalaw.org/lex//actview.xql?actid=No.%203%20of%202003"
            page = urllib.request.urlopen(quote_page)
            soup = BeautifulSoup(page, "lxml")
            Amending_Acts=soup.find("div", attrs={"id": "LongTitle_1_V7_t7"})
            print(Amending_Acts.text)

            time = soup.find("select", id="cboPIT")
            print("Amendment history.")
            print(time.text)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/A/Anti-Corruption%20and%20Economic%20Crimes%20Cap.%2065%20-%20No.%203%20of%202003/docs/Anti-CorruptionandEconomicCrimesAct3of2003.pdf"
            r = requests.get(Download, allow_redirects=True)
            toaster.show_toast("Downloaded Online Version ",
                                        " ",
                                        icon_path="scrapercon.ico",
                                        duration=5)
            open('Anti-corruption and Economic Crimes .pdf', 'wb').write(r.content)
            print("-"*50)
            hansard=input("Fetch Parliamentary Debates? [y/n] ").lower()
            if hansard=="y":
                hansard1= "https://books.google.co.ke/books?id=nMunxxOUWwEC&pg=PT26&lpg=PT26&dq=Kenya+assembly+hansard+anti-corruption+and+economic+crimes+Act&source=bl&ots=gkGpuAupKH&sig=kI4ZrNGOzYpJUNLv4swvmTxyBvs&hl=en&sa=X&ved=2ahUKEwiNq4fG-c3cAhUvxoUKHcaFB1MQ6AEwCXoECAAQAQ#v=onepage&q=Kenya%20assembly%20hansard%20anti-corruption%20and%20economic%20crimes%20Act&f=false"
                r = requests.get(hansard1, allow_redirects=True)
                hansard2= "https://books.google.co.ke/books?id=C5-GT140qskC&pg=PT2&lpg=PT2&dq=Kenya+assembly+hansard+anti-corruption+and+economic+crimes+Act&source=bl&ots=t_pTUmZPwy&sig=ZbcSXsAdq60_2ZgvSzIqvsZ6EMA&hl=en&sa=X&ved=2ahUKEwjI0KDc-s3cAhUM6RoKHbBCCno4ChDoATAAegQIBRAB#v=onepage&q=Kenya%20assembly%20hansard%20anti-corruption%20and%20economic%20crimes%20Act&f=false"
                r = requests.get(hansard2, allow_redirects=True)
                hansard3= "https://books.google.co.ke/books?id=T3OuOCEbA3sC&pg=PT24&lpg=PT24&dq=Kenya+assembly+hansard+anti-corruption+and+economic+crimes+Act&source=bl&ots=4ihfRGAqua&sig=RvEWpfezO80VU4Fv1_K7JVJhUWY&hl=en&sa=X&ved=2ahUKEwjI0KDc-s3cAhUM6RoKHbBCCno4ChDoATAFegQIAxAB#v=onepage&q=Kenya%20assembly%20hansard%20anti-corruption%20and%20economic%20crimes%20Act&f=false"
                r = requests.get(hansard3, allow_redirects=True)
                open('Ethics and Anti-corruption Hansard 23August2007.html', 'wb').write(r.content)
                open('Ethics and Anti-corruption Hansard 15April2008.html', 'wb').write(r.content)
                open('Ethics and Anti-corruption Hansard 15Sept2007.html', 'wb').write(r.content)
                toaster.show_toast("Downloads complete ",
                                        "Saved as *Ethics and Anit-corruption Hansards* ",
                                        icon_path="scrapercon.ico",
                                        duration=5)

    ###
        if actTitle=="Limitation of Actions" or actTitle=="statute of limitations" or actTitle=="limitations" or actTitle=="Limitations of Actions " or actTitle=="Limitations of Action " or actTitle=="limitation":
         
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%2022"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
           
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            quote_page = "http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%2022"
            page = urllib.request.urlopen(quote_page)
            soup = BeautifulSoup(page, "lxml")
            Amending_Acts=soup.find("div", attrs={"id": "longTitle_1_V4_t4"})
            print(Amending_Acts.text)

            time = soup.find("select", id="cboPIT")
            print("Amendment history")
            print(time.text)
            
            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/L/Limitation%20of%20Actions%20Act%20Cap.%2022%20-%20No.%2021%20of%201968/docs/LimitationofActionsAct21of1968.pdf"
            r = requests.get(Download, allow_redirects=True)
            toaster.show_toast("Downloaded Online Version ",
                                        " ",
                                        icon_path="scrapercon.ico",
                                        duration=5)
            open('Limitation of Actions .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Companies" or actTitle=="companies"  or actTitle=="companies act" or actTitle=="Companies Act" or actTitle=="Commercial":
                   
            print("Loading........")
            toaster.show_toast("Processing speed",
            "Some statutes are large in size and may be slow to process than others. ",
            icon_path="scrapercon.ico",
            duration=5)
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=No.%2017%20of%202015"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
           
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            quote_page = "http://www.kenyalaw.org/lex//actview.xql?actid=No.%2017%20of%202015"
            page = urllib.request.urlopen(quote_page)
            soup = BeautifulSoup(page, "lxml")
            Amending_Acts=soup.find("div", attrs={"id": "longTitle_1_V5_t5"})
            print(Amending_Acts.text)

            time = soup.find("select", id="cboPIT")
            print("Amendment history")
            print(time.text)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/C/Companies%20Act%20-%20No.%2017%20of%202015/docs/CompaniesAct17of2015.pdf"
            r = requests.get(Download, allow_redirects=True)
            toaster.show_toast("Downloaded Online Version ",
                                        " ",
                                        icon_path="scrapercon.ico",
                                        duration=5)
            open('Companies .pdf', 'wb').write(r.content)

            Download1=input ("Get the Companies (General) Regulations, 2015? [y/n] ").lower()
            if Download1=="yes" or Download1=="y":
                Download1= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/C/Companies%20Act%20-%20No.%2017%20of%202015/subsidiary%20legislation/docs/CompaniesAct17of2015_subsidiary.pdf"
                r = requests.get(Download1, allow_redirects=True)
                toaster.show_toast("Companies (General) Regulations ",
                                        "Download complete ",
                                        icon_path="scrapercon.ico",
                                        duration=5)
            open('Companies Regulations.pdf', 'wb').write(r.content)

     ###               
        if actTitle=="Anti-Counterfeit" or actTitle=="Anti-counterfeit" or actTitle=="anti-counterfeit" or actTitle=="counterfeit":
         
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=No.%2013%20of%202008"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
         
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            quote_page = "http://www.kenyalaw.org/lex//actview.xql?actid=No.%2013%20of%202008"
            page = urllib.request.urlopen(quote_page)
            soup = BeautifulSoup(page, "lxml")
            Amending_Acts=soup.find("div", attrs={"id": "longTitle_1_V4_t4"})
            print(Amending_Acts.text)

            time = soup.find("select", id="cboPIT")
            print("Amendment history")
            print(time.text)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/A/Anti-Counterfeit%20Act%20Cap.%20130A%20-%20No.%2013%20of%202008/docs/Anti-CounterfeitAct13of2008.pdf"
            r = requests.get(Download, allow_redirects=True)
            toaster.show_toast("Downloaded Online Version ",
                                        " ",
                                        icon_path="scrapercon.ico",
                                        duration=5)
            open('Anti-Counterfeit .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Anti-Doping" or actTitle=="Anti-doping" or actTitle=="anti-doping" or actTitle=="doping":
         
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=No.%205%20of%202016"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
        
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)
                
            quote_page = "http://www.kenyalaw.org/lex//actview.xql?actid=No.%205%20of%202016"
            page = urllib.request.urlopen(quote_page)
            soup = BeautifulSoup(page, "lxml")
            Amending_Acts=soup.find("div", attrs={"id": "longTitle_V1_t1"})
            print(Amending_Acts.text)

            time = soup.find("select", id="cboPIT")
            print("Amendment history")
            print(time.text)


            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/A/Anti-Doping%20Act%20-%20No.%205%20of%202016/docs/Anti-DopingAct5of2016.pdf"
            r = requests.get(Download, allow_redirects=True)
            toaster.show_toast("Downloaded Online Version ",
                                        " ",
                                        icon_path="scrapercon.ico",
                                        duration=5)
            open('Anti-Doping .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Appellate Jurisdiction" or actTitle=="appellate jurisdiction" or actTitle=="Jurisdiction" or actTitle=="jurisdiction" or actTitle=="Appellate" or actTitle=="appellate":
           
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%209"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
          
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)
                
            quote_page = "http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%209"
            page = urllib.request.urlopen(quote_page)
            soup = BeautifulSoup(page, "lxml")
            Amending_Acts=soup.find("div", attrs={"id": "longTitle_1_V3_t3"})
            print(Amending_Acts.text)

            time = soup.find("select", id="cboPIT")
            print("Amendment history")
            print(time.text)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/A/Appellate%20Jurisdiction%20Act%20Cap.%209%20-%20No.%2015%20of%201977/docs/AppellateJurisdictionAct15of1977.pdf"
            r = requests.get(Download, allow_redirects=True)
            toaster.show_toast("Downloaded Online Version ",
                                        " ",
                                        icon_path="scrapercon.ico",
                                        duration=5)
            open('Appellate Jurisdiction KLR.pdf', 'wb').write(r.content)

    ###
        if actTitle=="Arbitration" or actTitle=="arbitration":
           
            
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=No.%204%20of%201995"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
         
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            quote_page = "http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%209"
            page = urllib.request.urlopen(quote_page)
            soup = BeautifulSoup(page, "lxml")
            Amending_Acts=soup.find("div", attrs={"id": "longTitle_1_V3_t3"})
            print(Amending_Acts.text)

            time = soup.find("select", id="cboPIT")
            print("Amendment history")
            print(time.text)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/A/Arbitration%20Act%20Cap.%2049%20-%20No.%204%20of%201995/docs/ArbitrationAct4of1995.pdf"
            r = requests.get(Download, allow_redirects=True)
            toaster.show_toast("Downloaded Online Version ",
                                        " ",
                                        icon_path="scrapercon.ico",
                                        duration=5)
            open('Arbitration KLR.pdf', 'wb').write(r.content)

    ###
        if actTitle=="Architects and Quantity Surveyors" or actTitle=="architects and quantity Surveyors" or actTitle=="architects" or actTitle=="surveyors" or actTitle=="surveyor" or actTitle=="architect" or actTitle=="Architect" or actTitle=="Surveyor":
            
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%20525"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
          
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            quote_page = "http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%20525"
            page = urllib.request.urlopen(quote_page)
            soup = BeautifulSoup(page, "lxml")
            Amending_Acts=soup.find("div", attrs={"class": "act-long-title"})
            print(Amending_Acts.text)

##            time = soup.find("select", id="cboPIT")
##            print("Amendment history")
##            print(time.text)   

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/A/Architects%20and%20Quantity%20Surveyors%20Act%20Cap.%20525%20-/docs/ArchitectsandQuantitySurveyorsAct.pdf"
            r = requests.get(Download, allow_redirects=True)
            toaster.show_toast("Downloaded Online Version ",
                                        " ",
                                        icon_path="scrapercon.ico",
                                        duration=5)
            open('Architect KLR.pdf', 'wb').write(r.content)

    ###
        if actTitle=="Industrial Property Act" or actTitle=="industrial property act" or actTitle=="IPA" or actTitle=="ipa":
          
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=No.%203%20of%202001"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
          
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            quote_page = "http://www.kenyalaw.org/lex//actview.xql?actid=No.%203%20of%202001"
            page = urllib.request.urlopen(quote_page)
            soup = BeautifulSoup(page, "lxml")
            Amending_Acts=soup.find("div", attrs={"id": "longTitle_1_V4_t4"})
            print(Amending_Acts.text)

            time = soup.find("select", id="cboPIT")
            print("Amendment History")
            print(time.text)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/I/Industrial%20Property%20Act%20Cap.%20509%20-%20No.%203%20of%202001/docs/IndustrialPropertyAct3of2001.pdf"
            r = requests.get(Download, allow_redirects=True)
            toaster.show_toast("Downloaded Online Version ",
                                        " ",
                                        icon_path="scrapercon.ico",
                                        duration=5)
            open('Industrial Property KLR.pdf', 'wb').write(r.content)

    ###
        if actTitle=="Armed Forces(Out of bounds areas" or actTitle=="armed forces" or actTitle=="armed" or actTitle=="forces":
           
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%20202"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")

            
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            quote_page = "http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%20202"
            page = urllib.request.urlopen(quote_page)
            soup = BeautifulSoup(page, "lxml")
            Amending_Acts=soup.find("div", attrs={"class": "act-long-title"})
            print(Amending_Acts.text) 
            

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/A/Armed%20Forces%20Out%20of%20Bounds%20Areas%20Act%20Cap.%20202%20-%20No.%2024%20of%201948/docs/ArmedForcesOutofBoundsAreasAct.pdf"
            r = requests.get(Download, allow_redirects=True)
            toaster.show_toast("Downloaded Online Version ",
                                        " ",
                                        icon_path="scrapercon.ico",
                                        duration=5)
            open('Armed Forces KLR.pdf', 'wb').write(r.content)

    ###
        if actTitle=="Asian officers' family pensions" or actTitle=="asian officers' family pensions" or actTitle=="asian pension" or actTitle=="asian family pension" or actTitle=="asian family" or actTitle=="asian pension":
          
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%20194"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
           
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)
                
            quote_page = "http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%20194"
            page = urllib.request.urlopen(quote_page)
            soup = BeautifulSoup(page, "lxml")
            Amending_Acts=soup.find("div", attrs={"class": "act-long-title"})
            print(Amending_Acts.text)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/A/Asian%20Officers%20Family%20Pensions%20Act%20Cap.%20194%20-%20No.%2010%20of%201942/docs/AsianOfficersFamilyPensionsAct10of1942.pdf"
            r = requests.get(Download, allow_redirects=True)
            toaster.show_toast("Downloaded Online Version ",
                                        " ",
                                        icon_path="scrapercon.ico",
                                        duration=5)
            open('Asian_officers_family_pensions.pdf', 'wb').write(r.content)
            
    ###
        if actTitle=="Asian widows' and orphans' pensions" or actTitle=="asian widows and orphans pensions" or actTitle=="asian widows" or actTitle=="widows and orphans pension" or actTitle=="asian widow" or actTitle=="asian orphan":
           
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%20193"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
        
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)
                
            quote_page = "http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%20193"
            page = urllib.request.urlopen(quote_page)
            soup = BeautifulSoup(page, "lxml")
            Amending_Acts=soup.find("div", attrs={"class": "act-long-title"})
            print(Amending_Acts.text)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/A/Asian%20Officers%20Family%20Pensions%20Act%20Cap.%20194%20-%20No.%2010%20of%201942/docs/AsianOfficersFamilyPensionsAct10of1942.pdf"
            r = requests.get(Download, allow_redirects=True)
            toaster.show_toast("Downloaded Online Version ",
                                        " ",
                                        icon_path="scrapercon.ico",
                                        duration=5)
            
            open('Asian widows and orphans pensions KLR.pdf', 'wb').write(r.content)

    ###
        if actTitle=="Assumption of the office of president" or actTitle=="assumption of the office of president" or actTitle=="assumption of president" or actTitle=="assuming office" or actTitle=="assumption of office":
          
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=No.%2021%20of%202012"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
        
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            quote_page = "http://www.kenyalaw.org/lex//actview.xql?actid=No.%2021%20of%202012"
            page = urllib.request.urlopen(quote_page)
            soup = BeautifulSoup(page, "lxml")
            Amending_Acts=soup.find("div", attrs={"class": "act-long-title"})
            print(Amending_Acts.text)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/A/Assumption%20of%20the%20Office%20of%20President%20Act%20Cap.%20180%20-%20No.%2021%20of%202012/docs/AssumptionoftheOfficeofPresidentAct21of2012.pdf"
            r = requests.get(Download, allow_redirects=True)
            toaster.show_toast("Downloaded Online Version ",
                                        " ",
                                        icon_path="scrapercon.ico",
                                        duration=5)
            open('Assumption of office of president .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Auctioneers" or actTitle=="auctioneers" or actTitle=="auctioneers act" or actTitle=="Auctioneers Act" or actTitle=="Auctioneer" or actTitle=="auctioneer":          
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=No.%205%20of%201996"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
       
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            quote_page = "http://www.kenyalaw.org/lex//actview.xql?actid=No.%205%20of%201996"
            page = urllib.request.urlopen(quote_page)
            soup = BeautifulSoup(page, "lxml")
            Amending_Acts=soup.find("div", attrs={"id": "longTitle_1_V5_t5"})
            print(Amending_Acts.text)

            time = soup.find("select", id="cboPIT")
            print("Amendment History")
            print(time.text)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/A/Auctioneers%20Act%20Cap.%20526%20-%20No.%205%20of%201996/docs/AuctioneersAct5of1996.pdf"
            r = requests.get(Download, allow_redirects=True)
            toaster.show_toast("Downloaded Online Version ",
                                        " ",
                                        icon_path="scrapercon.ico",
                                        duration=5)
            open('Auctioneers .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Bank of Baroda" or actTitle=="Baroda Bank" or actTitle=="baroda" or actTitle=="baroda bank act":
           
           
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=No.%207%20of%201992"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
         
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)
                
            quote_page = "http://www.kenyalaw.org/lex//actview.xql?actid=No.%207%20of%201992"
            page = urllib.request.urlopen(quote_page)
            soup = BeautifulSoup(page, "lxml")
            Amending_Acts=soup.find("div", attrs={"class": "act-long-title"})
            print(Amending_Acts.text)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/B/Bank%20of%20Baroda%20Kenya%20Limited%20Act%20Cap.%20488B%20-%20No.%207%20of%201992/docs/BankofBarodaKenyaLimitedAct7of1992.pdf"
            r = requests.get(Download, allow_redirects=True)
            toaster.show_toast("Downloaded Online Version ",
                                        " ",
                                        icon_path="scrapercon.ico",
                                        duration=5)
            open('Baroda_Bank.pdf', 'wb').write(r.content)

    ###
        if actTitle=="Employment Act" or actTitle=="Employment" or actTitle=="employment":
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=No.%2011%20of%202007"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
        
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            quote_page = "http://www.kenyalaw.org/lex//actview.xql?actid=No.%2011%20of%202007"
            page = urllib.request.urlopen(quote_page)
            soup = BeautifulSoup(page, "lxml")
            Amending_Acts=soup.find("div", attrs={"id": "longTitle_1_V5_t5"})
            print(Amending_Acts.text)
            
            time = soup.find("select", id="cboPIT")
            print("Amendment History")
            print(time.text)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/E/Employment%20Act%20Cap.%20226%20-%20No.%2011%20of%202007/docs/EmploymentAct11of2007.pdf"
            r = requests.get(Download, allow_redirects=True)
            toaster.show_toast("Downloaded Online Version ",
                                        " ",
                                        icon_path="scrapercon.ico",
                                        duration=5)
            open('Employment.pdf', 'wb').write(r.content)

    ###
        if actTitle=="Succession Act" or actTitle=="law of succession act" or actTitle=="succession act" or actTitle=="succession" or actTitle=="Succession":
            
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%20160"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
            
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)
                
            quote_page = "http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%20160"
            page = urllib.request.urlopen(quote_page)
            soup = BeautifulSoup(page, "lxml")
            Amending_Acts=soup.find("div", attrs={"id": "longTitle_1_V4_t4"})
            print(Amending_Acts.text)
             
            time = soup.find("select", id="cboPIT")
            print("Amendment History")
            print(time.text)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/L/Law%20of%20Succession%20Act%20Cap.%20160%20-%20No.%2014%20of%201972/docs/LawofSuccessionAct14of1972.pdf"
            r = requests.get(Download, allow_redirects=True)
            toaster.show_toast("Downloaded Online Version ",
                                        " ",
                                        icon_path="scrapercon.ico",
                                        duration=5)
            open('Succession .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Barclays Bank of Kenya Limited" or actTitle=="Barclays Bank" or actTitle=="barclays bank" or actTitle=="barclays":
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=No.%2011%20of%201978"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
          
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            quote_page = "http://www.kenyalaw.org/lex//actview.xql?actid=No.%2011%20of%201978"
            page = urllib.request.urlopen(quote_page)
            soup = BeautifulSoup(page, "lxml")
            Amending_Acts=soup.find("div", attrs={"class": "act-long-title"})
            print(Amending_Acts.text)


            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/B/Barclays%20Bank%20of%20Kenya%20Limited%20Act%20Cap.%20488A%20-%20No.%2011%20of%201978/docs/BarclaysBankofKenyaLimited%20Act11of1978.pdf"
            r = requests.get(Download, allow_redirects=True)
            toaster.show_toast("Downloaded Online Version ",
                                        " ",
                                        icon_path="scrapercon.ico",
                                        duration=5)
            open('Barclays Bank .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Basic Education" or actTitle=="basic education" or actTitle=="education":
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=No.%2014%20of%202013"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
            
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            quote_page = "http://www.kenyalaw.org/lex//actview.xql?actid=No.%2014%20of%202013"
            page = urllib.request.urlopen(quote_page)
            soup = BeautifulSoup(page, "lxml")
            Amending_Acts=soup.find("div", attrs={"id": "LongTitle_1_V3_t3"})
            print(Amending_Acts.text)

            time = soup.find("select", id="cboPIT")
            print("Amendment History")
            print(time.text)
            
            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/B/Basic%20Education%20Act%20No.%2014%20of%202013/docs/BasicEducationActNo14of2013.pdf"
            r = requests.get(Download, allow_redirects=True)
            toaster.show_toast("Downloaded Online Version ",
                                        " ",
                                        icon_path="scrapercon.ico",
                                        duration=5)
            open('Basic_Education .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Betting Lotteries Gaming" or actTitle=="betting, lotteries and gaming" or actTitle=="betting" or actTitle=="gaming" or actTitle=="lottery" or actTitle=="lotteries" or actTitle=="Betting":
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%20131"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")

            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            quote_page = "http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%20131"
            page = urllib.request.urlopen(quote_page)
            soup = BeautifulSoup(page, "lxml")
            Amending_Acts=soup.find("div", attrs={"id": "longTitle_1_V6_t6"})
            print(Amending_Acts.text)

            time = soup.find("select", id="cboPIT")
            print("Amendment History")
            print(time.text)

            
            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/B/Betting%20Lotteries%20and%20Gaming%20Act%20Cap.%20131%20-%20No.%209%20of%201966/docs/BettingLotteriesandGamingAct9of1966.pdf"
            r = requests.get(Download, allow_redirects=True)
            toaster.show_toast("Downloaded Online Version ",
                                        " ",
                                        icon_path="scrapercon.ico",
                                        duration=5)
            open('Betting_Lotteries_Gaming .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Bills of Exchange" or actTitle=="bills of exchange" or actTitle=="bills" or actTitle=="exchange":
          
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%2027"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
            
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)
         
            quote_page = "http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%2027"
            page = urllib.request.urlopen(quote_page)
            soup = BeautifulSoup(page, "lxml")
            Amending_Acts=soup.find("div", attrs={"id": "longTitle_1_V6_t6"})
            print(Amending_Acts.text)

            time = soup.find("select", id="cboPIT")
            print("Amendment History")
            print(time.text)
            
            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/B/Bills%20of%20Exchange%20Act%20Cap.%2027%20-%20No.%207%20of%201927/docs/BillsofExchangeAct7of1927.pdf"
            r = requests.get(Download, allow_redirects=True)
            toaster.show_toast("Downloaded Online Version ",
                                        " ",
                                        icon_path="scrapercon.ico",
                                        duration=5)
            open('Bills of Exchange .pdf', 'wb').write(r.content)       
    ###
        if actTitle=="BIOSAFETY" or actTitle=="biosafety" or actTitle=="Biosafety Act":
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=No.%202%20of%202009"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
            
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)
            
            quote_page = "http://www.kenyalaw.org/lex//actview.xql?actid=No.%202%20of%202009"
            page = urllib.request.urlopen(quote_page)
            soup = BeautifulSoup(page, "lxml")
            Amending_Acts=soup.find("div", attrs={"class": "act-long-title"})
            print(Amending_Acts.text)

            time = soup.find("select", id="cboPIT")
            print("Amendment History")
            print(time.text)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/B/Biosafety%20Act%20Cap.%20321A%20-%20No.%202%20of%202009/docs/BiosafetyAct2of2009.pdf"
            r = requests.get(Download, allow_redirects=True)
            toaster.show_toast("Downloaded Online Version ",
                                        " ",
                                        icon_path="scrapercon.ico",
                                        duration=5)
            open('Biosafety .pdf', 'wb').write(r.content)

            Download1=input ("Do you need the Biosafety rules? [y/n] ").lower()
            if Download1=="yes" or Download1=="y":
                Download1= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/B/Biosafety%20Act%20Cap.%20321A%20-%20No.%202%20of%202009/subsidiary%20legislation/docs/BiosafetyAct2of2009_subsidiary.pdf"
                r = requests.get(Download1, allow_redirects=True)
                Download1= "Downloading Biosafety Rules as *biosafety rules*."
                list=textwrap3.wrap(Download1, width=90)
                for element in list:
                    print(element)
                print("----Download complete----.")
                open('Biosafety Rules.pdf', 'wb').write(r.content)
    ###
        if actTitle=="Births and Deaths Registration" or actTitle=="births and deaths registration" or actTitle=="births" or actTitle=="deaths":
           
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%20149"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
            
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)
            
            quote_page = "http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%20149"
            page = urllib.request.urlopen(quote_page)
            soup = BeautifulSoup(page, "lxml")
            Amending_Acts=soup.find("div", attrs={"class": "act-long-title"})
            print(Amending_Acts.text)

##            time = soup.find("select", id="cboPIT")
##            print("Amendment History")
##            print(time.text)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/B/Births%20and%20Deaths%20Registration%20Act%20Cap.%20149%20-%20Act%20No.%202%20of%201928/docs/BirthsandDeathsRegistrationAct2of1928.pdf"
            r = requests.get(Download, allow_redirects=True)
            toaster.show_toast("Downloaded Online Version ",
                                        " ",
                                        icon_path="scrapercon.ico",
                                        duration=5)
            open('Births and Deaths Registration .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Books and Newspapers" or actTitle=="books and newspapers" or actTitle=="books" or actTitle=="newspapers":
           
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%20111"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
            
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)
            
            quote_page = "http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%20111"
            page = urllib.request.urlopen(quote_page)
            soup = BeautifulSoup(page, "lxml")
            Amending_Acts=soup.find("div", attrs={"id": "longTitle_1_V2_t2"})
            print(Amending_Acts.text)

            time = soup.find("select", id="cboPIT")
            print("Amendment History")
            print(time.text)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/B/Books%20and%20Newspapers%20Act%20Cap.%20111%20-%20No.%2027%20of%201960/docs/BooksandNewspapersAct32of1930.pdf"
            r = requests.get(Download, allow_redirects=True)
            toaster.show_toast("Downloaded Online Version ",
                                        " ",
                                        icon_path="scrapercon.ico",
                                        duration=5)
            open('Books and Newspapers .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Borstal Institutions" or actTitle=="borstal institutions" or actTitle=="borstal":
                
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%2092"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
            
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)
            
            quote_page = "http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%2092"
            page = urllib.request.urlopen(quote_page)
            soup = BeautifulSoup(page, "lxml")
            Amending_Acts=soup.find("div", attrs={"class": "act-long-title"})
            print(Amending_Acts.text)

##            time = soup.find("select", id="cboPIT")
##            print("Amendment History")
##            print(time.text)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/B/Borstal%20Institutions%20Act%20Cap.%2092%20-%20No.%2023%20of%201963/docs/BorstalInstitutionsAct23of1963.pdf"
            r = requests.get(Download, allow_redirects=True)
            toaster.show_toast("Downloaded Online Version ",
                                        " ",
                                        icon_path="scrapercon.ico",
                                        duration=5)
            open('Borstal Institutions .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Branding of stock" or actTitle=="branding" or actTitle=="branding of stock" or actTitle=="stock branding":
            
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%20357"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
            
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            quote_page = "http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%20357"
            page = urllib.request.urlopen(quote_page)
            soup = BeautifulSoup(page, "lxml")
            Amending_Acts=soup.find("div", attrs={"class": "act-long-title"})
            print(Amending_Acts.text)

            time = soup.find("select", id="cboPIT")
            print("Amendment History")
            print(time.text)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/B/Branding%20of%20Stock%20Act%20Cap.%20357%20-%20No.%2012%20of%201907/docs/BrandingofStockAct12of1907.pdf"
            r = requests.get(Download, allow_redirects=True)
            toaster.show_toast("Downloaded Online Version ",
                                        " ",
                                        icon_path="scrapercon.ico",
                                        duration=5)
            open('Branding of stock .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Breast Milk Substitutes" or actTitle=="breast milk substitutes" or actTitle=="Breast milk" or actTitle=="breast milk" or actTitle=="breast":
                 
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=No.%2034%20of%202012"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
            
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)
             
            quote_page = "http://www.kenyalaw.org/lex//actview.xql?actid=No.%2034%20of%202012"
            page = urllib.request.urlopen(quote_page)
            soup = BeautifulSoup(page, "lxml")
            Amending_Acts=soup.find("div", attrs={"id": "longTitle_1"})
            print(Amending_Acts.text)

##            time = soup.find("select", id="cboPIT")
##            print("Amendment History")
##            print(time.text)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/B/Breast%20Milk%20Substitutes%20Regulation%20and%20Control%20Act%20No.%2034%20of%202012/docs/BreastMilkSubstitutesRegulationandControlAct34of2012.pdf"
            r = requests.get(Download, allow_redirects=True)
            toaster.show_toast("Downloaded Online Version ",
                                        " ",
                                        icon_path="scrapercon.ico",
                                        duration=5)
            open('Breast Milk .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Bretton Woods Agreements" or actTitle=="bretton woods" or actTitle=="Bretton Woods" or actTitle=="bretton woods" or actTitle=="bretton" or actTitle=="Bretton":
                
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%20464"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
            
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)
            
            quote_page = "http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%20464"
            page = urllib.request.urlopen(quote_page)
            soup = BeautifulSoup(page, "lxml")
            Amending_Acts=soup.find("div", attrs={"class": "act-long-title"})
            print(Amending_Acts.text)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/B/Bretton%20Woods%20Agreement%20Act%20Cap.%20464%20-%20No.%2038%20of%201963/docs/BrettonWoodsAgreementAct38of1963.pdf"
            r = requests.get(Download, allow_redirects=True)
            toaster.show_toast("Downloaded Online Version ",
                                        " ",
                                        icon_path="scrapercon.ico",
                                        duration=5)
            open('Bretton Woods Agreements .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Bribery Act" or actTitle=="bribery act" or actTitle=="Bribery" or actTitle=="bribery":

            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=No.%2047%20of%202016"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
            
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)
               
            quote_page = "http://www.kenyalaw.org/lex//actview.xql?actid=No.%2047%20of%202016"
            page = urllib.request.urlopen(quote_page)
            soup = BeautifulSoup(page, "lxml")
            Amending_Acts=soup.find("div", attrs={"class": "act-long-title"})
            print(Amending_Acts.text)
            
            time = soup.find("select", id="cboPIT")
            print("Amendment History")
            print(time.text)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/B/Bribery%20Act%20No.%2047%20of%202016/docs/BriberyAct47of2016.pdf"
            r = requests.get(Download, allow_redirects=True)
            toaster.show_toast("Downloaded Online Version ",
                                        " ",
                                        icon_path="scrapercon.ico",
                                        duration=5)
            open('Bribery .pdf', 'wb').write(r.content)

    ###
        if actTitle=="British Standard Portland Cement Company Limited" or actTitle=="Bamburi Factory" or actTitle=="Bamburi" or actTitle=="bamburi" or actTitle== "bamburi cement":
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%20515"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
            
            title= soup.find("div", attrs={"class": "act-title"})
            for title in title:
                print(title.string)

            quote_page = "http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%20515"
            page = urllib.request.urlopen(quote_page)
            soup = BeautifulSoup(page, "lxml")
            Amending_Acts=soup.find("div", attrs={"class": "act-long-title"})
            print(Amending_Acts.text)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/B/British%20Std%20Portland%20Cement%20Cpy%20Ltd%20Bamburi%20FactoryActCap.515-No.61of1951/docs/BritishStandardPortlandCementCompanyLtdBamburiFactoryAct61of1951.pdf"
            r = requests.get(Download, allow_redirects=True)
            toaster.show_toast("Downloaded Online Version ",
                                        " ",
                                        icon_path="scrapercon.ico",
                                        duration=5)
            open('British Standard Portland Cement Company Limited (Bamburi) .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Brokers" or actTitle=="brokers" or actTitle=="Broker" or actTitle=="broker":
           
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%20527"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
            
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)
            
            quote_page = "http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%20515"
            page = urllib.request.urlopen(quote_page)
            soup = BeautifulSoup(page, "lxml")
            Amending_Acts=soup.find("div", attrs={"id": "longTitle_1_V2_t2"})
            print(Amending_Acts.text)
            
            time = soup.find("select", id="cboPIT")
            print("Amendment History")
            print(time.text)
            
            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/B/Brokers%20Act%20Cap.%20527%20-%20No.%2056%20of%201930/docs/BrokersAct56of1930.pdf"
            r = requests.get(Download, allow_redirects=True)
            toaster.show_toast("Downloaded Online Version ",
                                        " ",
                                        icon_path="scrapercon.ico",
                                        duration=5)
            open('Brokers .pdf', 'wb').write(r.content)

     ###       
        if actTitle=="Building Societies" or actTitle=="building societies" or actTitle=="building" or actTitle=="Building":
       
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%20489"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
            
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)
            
            quote_page = "http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%20489"
            page = urllib.request.urlopen(quote_page)
            soup = BeautifulSoup(page, "lxml")
            Amending_Acts=soup.find("div", attrs={"id": "longTitle_1_V9_t9"})
            print(Amending_Acts.text)
            
            time = soup.find("select", id="cboPIT")
            print("Amendment History")
            print(time.text)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/B/Building%20Societies%20Act%20Cap.%20489%20-%20No.%2029%20of%201956/docs/BuildingSocietiesAct29of1956.pdf"
            r = requests.get(Download, allow_redirects=True)
            toaster.show_toast("Downloaded Online Version ",
                                        " ",
                                        icon_path="scrapercon.ico",
                                        duration=5)
            open('Building Societies .pdf', 'wb').write(r.content)

     ###       
        if actTitle=="Bukura Agricultural College" or actTitle=="Bukura College" or actTitle=="bukura college" or actTitle=="bukura" or actTitle=="Bukura":
        
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=No.%205%20of%201999"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")

            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)
            
            quote_page = "http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%20489"
            page = urllib.request.urlopen(quote_page)
            soup = BeautifulSoup(page, "lxml")
            Amending_Acts=soup.find("div", attrs={"id": "longTitle_1_V3_t3"})
            print(Amending_Acts.text)
            
            time = soup.find("select", id="cboPIT")
            print("Amendment History.")
            print(time.text)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/B/Bukura%20Agricultural%20College%20Act%20Cap.%20348%20-%20No.%205%20of%201999/docs/BukuraAgriculturalCollegeAct5of1999.pdf"
            r = requests.get(Download, allow_redirects=True)
            toaster.show_toast("Downloaded Online Version ",
                                        " ",
                                        icon_path="scrapercon.ico",
                                        duration=5)
            open('Bukura Agricultural College .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Business Registration Service" or actTitle=="business registration" or actTitle=="business registration service" or actTitle=="BRS" or actTitle=="brs" or actTitle=="Business Registration":
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=No.%2015%20of%202015"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
            
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            quote_page = "http://www.kenyalaw.org/lex//actview.xql?actid=No.%2015%20of%202015"
            page = urllib.request.urlopen(quote_page)
            soup = BeautifulSoup(page, "lxml")
            Amending_Acts=soup.find("div", attrs={"id": "longTitle_1"})
            print(Amending_Acts.text)
            
            time = soup.find("select", id="cboPIT")
            print("Amendment History.")
            print(time.text)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/B/Business%20Registration%20Act%20No.%2015%20of%202015/docs/BusinessRegistrationServiceAct15of2015.pdf"
            r = requests.get(Download, allow_redirects=True)
            toaster.show_toast("Downloaded Online Version ",
                                        " ",
                                        icon_path="scrapercon.ico",
                                        duration=5)
            open('Business Registration Service .pdf', 'wb').write(r.content)

    ###
        if actTitle=="CABINET SECRETARY TO THE TREASURY (INCORPORATION)" or actTitle=="CS Treasury" or actTitle=="Cabinet Secretary to the Treasury (Incorporation)" or actTitle=="Cabinet" or actTitle=="cabinet" or actTitle=="Secretary" or actTitle=="secretary" or actTitle=="cabinet secretary" or actTitle=="Cabinet secretary" or actTitle=="Cabinet Secretary" or actTitle=="treasury" or actTitle=="Treasury":
                   
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%20101"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
            
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)
            
            quote_page = "http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%20101"
            page = urllib.request.urlopen(quote_page)
            soup = BeautifulSoup(page, "lxml")
            Amending_Acts=soup.find("div", attrs={"id": "longTitle_1_V2_t2"})
            print(Amending_Acts.text)

            time = soup.find("select", id="cboPIT")
            print("Amendment History.")
            print(time.text)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/C/Cabinet%20Secretary%20to%20the%20Treasury%20Incorporation%20Act%20Cap.%20101%20-%20No.%202%20of%201962/docs/CabinetSecretarytotheTreasuryIncorporationAct2of1962.pdf"
            r = requests.get(Download, allow_redirects=True)
            toaster.show_toast("Downloaded Online Version ",
                                        " ",
                                        icon_path="scrapercon.ico",
                                        duration=5)
            open('Cabinet Secretary to the Treasury (Incorporation) .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Cancer Prevention and Control" or actTitle=="Cancer prevention" or actTitle=="cancer prevention" or actTitle=="Cancer" or actTitle=="cancer" or actTitle=="Cancer Control" or actTitle=="cancer control" or actTitle=="Cancer Prevention":
           
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=No.%2015%20of%202012"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
            
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)
            
            quote_page = "http://www.kenyalaw.org/lex//actview.xql?actid=No.%2015%20of%202012"
            page = urllib.request.urlopen(quote_page)
            soup = BeautifulSoup(page, "lxml")
            Amending_Acts=soup.find("div", attrs={"class": "act-long-title"})
            print(Amending_Acts.text)

            time = soup.find("select", id="cboPIT")
            print("Amendment History.")
            print(time.text)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/C/Cancer%20Prevention%20and%20Control%20Act%20Cap.%20246B%20-%20No.%2015%20of%202012/docs/CancerPreventionandControlAct15of2012.pdf"
            r = requests.get(Download, allow_redirects=True)
            toaster.show_toast("Downloaded Online Version ",
                                        " ",
                                        icon_path="scrapercon.ico",
                                        duration=5)
            open('Cancer Prevention and Control .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Capital Markets" or actTitle=="capital markets" or actTitle=="Capital" or actTitle=="capital" or actTitle=="Markets" or actTitle=="markets" or actTitle=="cma" or actTitle=="CMA":
              
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%20485A"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
            
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)
            
            quote_page = "http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%20485A"
            page = urllib.request.urlopen(quote_page)
            soup = BeautifulSoup(page, "lxml")
            Amending_Acts=soup.find("div", attrs={"id":"longTitle_1_V19_t19"})
            print(Amending_Acts.text)
                
            print("\n")
            time = soup.find("select", id="cboPIT")
            print("Amendment History")
            print(time.text)      

            Download= "https://www.cma.or.ke/index.php?option=com_phocadownload&view=category&download=321:the-capital-markets-act-amended-2016&id=12:acts&Itemid=192"
            r = requests.get(Download, allow_redirects=True)
            toaster.show_toast("Downloaded Online Version ",
                                        " ",
                                        icon_path="scrapercon.ico",
                                        duration=5)
            open('Capital Markets eCMA.pdf', 'wb').write(r.content)
            print("***********************************************")

            masterPlan=input ("Would you like a copy of the Capital Markets Masterplan (2014-2023)? [y/n] ").lower()
            if masterPlan=="yes" or masterPlan=="y":
                print("Downloading Capital Markets Masterplan (2014-2023) as *CMA Masterplan*....")
                cma3= "http://cmmp.or.ke/index.php"
                r = requests.get(cma3, allow_redirects=True)
                
                open('CMA Master plan eCMA.html', 'wb').write(r.content)
           
            cma3=input ("Would you like a copy of the Capital Markets subsidiary legislation? [y/n] ").lower()
            if cma3=="yes" or cma3=="y":
                cma3= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/C/Capital%20Markets%20Act%20Cap.%20485A%20Act%20No.%2017%20of%201989/subsidiary%20legislation/docs/Capital%20MarketsAct17of1989_subsidiary.pdf"
                r = requests.get(cma3, allow_redirects=True)
                toaster.show_toast("Downloads complete. check folder for the Masterplan and regulations. ",
                                        " ",
                                        icon_path="scrapercon.ico",
                                        duration=5)
            print("-"*50)
            open('CMA Regulations.pdf', 'wb').write(r.content)

    ###
        if actTitle=="Carriage By Air" or actTitle=="carriage by air" or actTitle=="Carriage" or actTitle=="carriage" or actTitle=="Air" or actTitle=="Carriage by air":
           
        
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=No.%202%20of%201993"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
            
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)
             
            quote_page = "http://www.kenyalaw.org/lex//actview.xql?actid=No.%202%20of%201993"
            page = urllib.request.urlopen(quote_page)
            soup = BeautifulSoup(page, "lxml")
            Amending_Acts=soup.find("div", attrs={"class":"act-long-title"})
            print(Amending_Acts.text)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/C/Carriage%20by%20Air%20Act%20Cap.%20394A%20-%20No.%202%20of%201993/docs/CarriagebyAirAct2of1993.pdf"
            r = requests.get(Download, allow_redirects=True)
            toaster.show_toast("Downloaded Online Version ",
                                        " ",
                                        icon_path="scrapercon.ico",
                                        duration=5)
            
            open('Carriage by air .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Carriage of goods by sea" or actTitle=="Carriage of Goods By Sea" or actTitle=="Carriage of Goods by Sea" or actTitle=="carriage of goods by sea" or actTitle=="Carriage of Goods" or actTitle=="Carriage Goods" or actTitle=="carriage goods" or actTitle== "Carriage by sea" or actTitle== "carriage by sea":
          
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%20392"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")

            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)
            
            quote_page = "http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%20392"
            page = urllib.request.urlopen(quote_page)
            soup = BeautifulSoup(page, "lxml")
            Amending_Acts=soup.find("div", attrs={"class":"act-long-title"})
            print(Amending_Acts.text)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/C/Carriage%20of%20Goods%20by%20Sea%20Act%20Cap.%20392%20-%20No.%2016%20of%201926/docs/CarriageofGoodsbySeaAct16of1926.pdf"
            r = requests.get(Download, allow_redirects=True)
            toaster.show_toast("Downloaded Online Version ",
                                        " ",
                                        icon_path="scrapercon.ico",
                                        duration=5)
            open('Carriage by goods by sea .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Cattle Cleansing" or actTitle=="cattle cleansing" or actTitle=="Cattle" or actTitle=="Cleansing" or actTitle=="cattle" or actTitle=="cleansing":
            
            actTitle1= ["http://www.kenyalaw.org/lex//sublegview.xql?subleg=CAP.%20358"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
            
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)
            
            quote_page = "http://www.kenyalaw.org/lex//sublegview.xql?subleg=CAP.%20358"
            page = urllib.request.urlopen(quote_page)
            soup = BeautifulSoup(page, "lxml")
            Amending_Acts=soup.find("div", attrs={"class":"act-long-title"})
            print(Amending_Acts.text)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/C/Cattle%20Cleansing%20Act%20Cap.%20358%20-%20Act%20No.%2032%20of%201929/docs/CattleCleansingAct32of1929.pdf"
            r = requests.get(Download, allow_redirects=True)
            toaster.show_toast("Downloaded Online Version ",
                                        " ",
                                        icon_path="scrapercon.ico",
                                        duration=5)
            open('Cattle Cleansing .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Central Bank of Kenya" or actTitle=="central bank of Kenya" or actTitle=="Central Bank" or actTitle=="central bank" or actTitle=="cbk" or actTitle=="CBK":
           
             
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%20491"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")

            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)
            
            quote_page = "http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%20491"
            page = urllib.request.urlopen(quote_page)
            soup = BeautifulSoup(page, "lxml")
            Amending_Acts=soup.find("div", attrs={"id":"longTitle_1_V16_t21"})
            print(Amending_Acts.text)

            print("\n")
            time = soup.find("select", id="cboPIT")
            print("Amendment History")
            print(time.text) 

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/C/Central%20Bank%20of%20Kenya%20Act%20Cap.%20491%20-%20No.%2015%20of%201966/docs/CentralBankofKenyaAct15of1966.pdf"
            r = requests.get(Download, allow_redirects=True)
            toaster.show_toast("Downloaded Online Version ",
                                        " ",
                                        icon_path="scrapercon.ico",
                                        duration=5)
            open('Central Bank .pdf', 'wb').write(r.content)

            cbk3=input ("Do you need the Central Bank subsidiary legislation? [y/n] ").lower()
            if cbk3=="yes" or cbk3=="y":
                cbk3= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/C/Central%20Bank%20of%20Kenya%20Act%20Cap.%20491%20-%20No.%2015%20of%201966/subsidiary%20legislation/docs/CentralBankofKenyaAct15of1966_subsidiary.pdf"
                r = requests.get(cbk3, allow_redirects=True)
            open('CBK Regulations.pdf', 'wb').write(r.content)
            print(title.string, "regulations downloaded. Check folder. ")

    ###
        if actTitle=="Central Depositories" or actTitle=="central depositories" or actTitle=="Depositories" or actTitle=="depositories" or actTitle=="cda" or actTitle=="CDA":
            
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=No.%204%20of%202000"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
            
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)
     
            quote_page = "http://www.kenyalaw.org/lex//actview.xql?actid=No.%204%20of%202000"
            page = urllib.request.urlopen(quote_page)
            soup = BeautifulSoup(page, "lxml")
            Amending_Acts=soup.find("div", attrs={"id":"longTitle_1_V3_t3"})
            print(Amending_Acts.text)

            print("\n")
            time = soup.find("select", id="cboPIT")
            print("Amendment History")
            print(time.text)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/C/Central%20Depositories%20Act%20Cap.%20485C%20-%20No.%204%20of%202000/docs/CentralDepositoriesAct4of2000.pdf"
            r = requests.get(Download, allow_redirects=True)
            toaster.show_toast("Downloaded Online Version ",
                                        " ",
                                        icon_path="scrapercon.ico",
                                        duration=5)
            open('Central Depositories .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Certified Public Secretaries of Kenya" or actTitle=="certified public secretaries of Kenya" or actTitle=="CPS" or actTitle=="cps" or actTitle=="cpsk" or actTitle=="public secretaries":
           
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%20534"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
            
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)
            
            quote_page = "http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%20534"
            page = urllib.request.urlopen(quote_page)
            soup = BeautifulSoup(page, "lxml")
            Amending_Acts=soup.find("div", attrs={"id":"longTitle_1_V8_t8"})
            print(Amending_Acts.text)

            print("\n")
            time = soup.find("select", id="cboPIT")
            print("Amendment History")
            print(time.text)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/C/Certified%20Public%20Secretaries%20of%20Kenya%20Act%20Cap.%20534%20-%20No.%2012%20of%201988/docs/CertifiedPublicSecretariesofKenyaAct12of1988.pdf"
            r = requests.get(Download, allow_redirects=True)
            toaster.show_toast("Downloaded Online Version ",
                                        " ",
                                        icon_path="scrapercon.ico",
                                        duration=5)
            open('Certified Public Secretaries .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Public Procurement and Asset Disposal Act" or actTitle=="public procurement" or actTitle=="Procurement" or actTitle=="procurement" or actTitle=="ppda" or actTitle=="PPDA" or actTitle=="Public Procurement" : 
           
            actTitle1= ["http://kenyalaw.org/lex//actview.xql?actid=No.%2033%20of%202015"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
            
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)
            
            quote_page = "http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%20534"
            page = urllib.request.urlopen(quote_page)
            soup = BeautifulSoup(page, "lxml")
            Amending_Acts=soup.find("div", attrs={"id":"longTitle_1_V8_t8"})
            print(Amending_Acts.text)

            print("\n")
            time = soup.find("select", id="cboPIT")
            print("Amendment History")
            print(time.text)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/P/Public%20Procurement%20and%20Asset%20Disposal%20Act%20-%20No.%2033%20of%202015/docs/PublicProcurementAndAssetDisposalAct33of2015.pdf"
            r = requests.get(Download, allow_redirects=True)
            toaster.show_toast("Downloaded Online Version ",
                                        " ",
                                        icon_path="scrapercon.ico",
                                        duration=5)
            open('Public Procurement and Asset Disposal .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Civil Aviation" or actTitle=="civil aviation" or actTitle=="aviation":  
            
            actTitle1= ["http://kenyalaw.org/lex//actview.xql?actid=No.%2021%20of%202013"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
         
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            quote_page = "http://kenyalaw.org/lex//actview.xql?actid=No.%2021%20of%202013"
            page = urllib.request.urlopen(quote_page)
            soup = BeautifulSoup(page, "lxml")
            Amending_Acts=soup.find("div", attrs={"id":"longTitle_1_V3_t3"})
            print(Amending_Acts.text)

            print("\n")
            time = soup.find("select", id="cboPIT")
            print("Amendment History")
            print(time.text)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/C/Civil%20Aviation%20Act%20No.%2021%20of%202013/docs/CivilAviationActNo.21of2013.pdf"
            r = requests.get(Download, allow_redirects=True)
            toaster.show_toast("Downloaded Online Version ",
                                        " ",
                                        icon_path="scrapercon.ico",
                                        duration=5)
            open('Civil Aviation .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Entertainment Tax" or actTitle=="entertainment tax" or actTitle=="Entertainment" or actTitle=="entertainment":
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%20479"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")

            
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            quote_page = "http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%20479"
            page = urllib.request.urlopen(quote_page)
            soup = BeautifulSoup(page, "lxml")
            Amending_Acts=soup.find("div", attrs={"id":"longTitle_1_V2_t2"})
            print(Amending_Acts.text)

            print("\n")
            time = soup.find("select", id="cboPIT")
            print("Amendment History")
            print(time.text)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/E/Entertainment%20Tax%20Act%20Cap.%20479%20-%20No.%2063%20of%201950/docs/EntertainmentTaxAct63of1950.pdf"
            r = requests.get(Download, allow_redirects=True)
            toaster.show_toast("Downloaded Online Version ",
                                        " ",
                                        icon_path="scrapercon.ico",
                                        duration=5)
            open('Entertainment Tax .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Income Tax" or actTitle=="income tax" or actTitle=="Income tax" or actTitle=="Income" or actTitle=="income":
            
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%20470"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")

            
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)
                    
            quote_page = "http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%20470"
            page = urllib.request.urlopen(quote_page)
            soup = BeautifulSoup(page, "lxml")
            Amending_Acts=soup.find("div", attrs={"id":"longTitle_1_V11_t11"})
            print(Amending_Acts.text)

            print("\n")
            time = soup.find("select", id="cboPIT")
            print("Amendment History")
            print(time.text)

            print("\n")
            FinanceBill18=input("Would you like to view an analysis of the Finance Bill 2018? [y/n] ").lower()
            if FinanceBill18=="yes" or FinanceBill18=="y":
                FinanceBill18="https://home.kpmg.com/content/dam/kpmg/ke/pdf/tax/Finance%20Bill%202018%20Analysis.pdf"
                r = requests.get(FinanceBill18, allow_redirects=True)
                list=textwrap3.wrap(FinanceBill18, width=90)
                for element in list:
                    print(element)
                open('Finance Bill 2018-KPMG Analysis.pdf', 'wb').write(r.content)
                print("Download complete. ")
            print("\n")
            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/I/Income%20Tax%20Act%20Cap.%20470%20-%20No.%2016%20of%201973/docs/IncomeTaxAct2of1975.pdf"
            r = requests.get(Download, allow_redirects=True)
            toaster.show_toast("Downloaded Online Version ",
                                        " ",
                                        icon_path="scrapercon.ico",
                                        duration=5)
            open('Income Tax .pdf', 'wb').write(r.content)

            IncomeTaxBill18= "Hey, according to the Treasury website, there exists a draft Income Tax Bill, 2018. Donwloading it now.. "
            list=textwrap3.wrap(IncomeTaxBill18, width=90)
            for element in list:
                print(element)
            Download= "http://www.treasury.go.ke/tax/finance-bills.html?download=781:draft-income-tax-bill-2018"
            r = requests.get(Download, allow_redirects=True)
            open('Income Tax .pdf', 'wb').write(r.content)
            print("Download complete. Check folder for details. ")

    ###
        if actTitle=="VAT" or actTitle=="vat":
             
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=No.%2035%20of%202013"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")

            
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            quote_page = "http://www.kenyalaw.org/lex//actview.xql?actid=No.%2035%20of%202013"
            page = urllib.request.urlopen(quote_page)
            soup = BeautifulSoup(page, "lxml")
            Amending_Acts=soup.find("div", attrs={"id":"longTitle_1_V2_t2"})
            print(Amending_Acts.text)

            print("\n")
            time = soup.find("select", id="cboPIT")
            print("Amendment History")
            print(time.text)
            
            FinanceBill18=input("Would you like to view an analysis of the Finance Bill 2018? [y/n] ").lower()
            if FinanceBill18=="yes" or FinanceBill18=="y":
                FinanceBill18="https://home.kpmg.com/content/dam/kpmg/ke/pdf/tax/Finance%20Bill%202018%20Analysis.pdf"
                r = requests.get(FinanceBill18, allow_redirects=True)
                FinanceBill18= "Downloading analysis of the Finance Bill 2018 as *Finance Bill 2018-KPMG Analysis*."
                list=textwrap3.wrap(FinanceBill18, width=90)
                for element in list:
                    print(element)
                print("----Download complete----.")
            print("************************************************************")
            print("\n")
            open('Finance Bill 2018-KPMG Analysis.pdf', 'wb').write(r.content)
                
            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/V/Value%20Added%20Tax%20Act%20No.%2035%20of%202013/docs/ValueAddedTax_ActNo35of2013.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("I downloaded a copy of the VAT Act just in case. .")
            print("\n")
            open('Value Added Tax .pdf', 'wb').write(r.content)

            Download1=input ("Do you need the VAT subsidiary legislation? [y/n] ").lower()
            if Download1=="yes" or Download1=="y":
                Download1= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/V/Value%20Added%20Tax%20Act%20No.%2035%20of%202013/subsidiary%20legislation/docs/ValueAddedTaxAct7of1989_subsidairy.pdf"
                r = requests.get(Download1, allow_redirects=True)
                Download1= "Downloading VAT subsidiary legislation as *VAT Regulations*."
                list=textwrap3.wrap(Download1, width=90)
                for element in list:
                    print(element)
                print("----Download complete----.")
                open('VAT Regulations.pdf', 'wb').write(r.content)
          
    ###
        if actTitle=="Tax Appeals Tribunal" or actTitle=="tax appeals tribunal" or actTitle=="tax appeals" or actTitle=="Tax Appeals":
            
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=No.%2040%20of%202013"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")

            
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)
            print("\n")
            #time = soup.find("select", id="cboPIT")
            #print("Amendment History")
            #print(time.text)  
            print(User_Name, ", this is the time stamp of the legislation that matches your keyword: ")
            tat2= "Statute quarantined "
            list=textwrap3.wrap(tat2, width=95)
            for element in list:
                print(element)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/T/Tax%20Appeals%20Tribunal%20Act%20-%20No.%2040%20of%202013/docs/TaxAppealsTribunalActNo.40of2013.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("-----Downloaded Online Version-----.")
            print("\n")
            
            open('Tax Appeals Tribunal .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Excise Duty" or actTitle=="excise duty" or actTitle=="excise" or actTitle=="excise tax":
            
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=No.%2023%20of%202015"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")

            
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)
                    
            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)
            print("\n")
            #time = soup.find("select", id="cboPIT")
            #print("Amendment History")
            #print(time.text)  
            print(User_Name, ", this is the time stamp of the legislation that matches your keyword: ")
            ed2= "Last recorded amendments: See Finance Bill (National Assembly Bills 20) 2018. "
            list=textwrap3.wrap(ed2, width=95)
            for element in list:
                print(element)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/E/Excise%20Duty%20Act%20-%20No.%2023%20of%202015/docs/ExciseDutyAct23of2015.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("-----Downloaded Online Version-----.")
            print("\n")
            open('Excise Duty .pdf', 'wb').write(r.content)

            print("************************************************************")
            print("\n")
            FinanceBill18=input("Would you like to view an analysis of the Finance Bill 2018? [y/n] ").lower()
            if FinanceBill18=="yes" or FinanceBill18=="y":
                FinanceBill18="https://home.kpmg.com/content/dam/kpmg/ke/pdf/tax/Finance%20Bill%202018%20Analysis.pdf"
                r = requests.get(FinanceBill18, allow_redirects=True)
                FinanceBill18= "Downloading analysis of the Finance Bill 2018 as *Finance Bill 2018-KPMG Analysis*."
                list=textwrap3.wrap(FinanceBill18, width=90)
                for element in list:
                    print(element)
                print("----Download complete----.")
            print("************************************************************")
            print("\n")
            open('Finance Bill 2018-KPMG Analysis.pdf', 'wb').write(r.content)

    ###
        if actTitle=="Tax Procedures" or actTitle=="tax procedures":
            
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=No.%2029%20of%202015"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")

            
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)
                   
            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)
            print("\n")
            #time = soup.find("select", id="cboPIT")
            #print("Amendment History")
            #print(time.text)  
            print(User_Name, ", this is the time stamp of the legislation that matches your keyword: ")
            tpa2= "Last recorded amendments: see Finance Bill (National Assembly Bills 20) 2018. "
            list=textwrap3.wrap(tpa2, width=95)
            for element in list:
                print(element)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/T/Tax%20Procedures%20Act%20-%20No%2029%20of%202015/docs/TaxProceduresAct29of2015.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("-----Downloaded Online Version-----.")
            print("\n")
            open('Tax Procedures .pdf', 'wb').write(r.content)

            print("************************************************************")
            print("\n")
            FinanceBill18=input("Would you like to view an analysis of the Finance Bill 2018? [y/n] ").lower()
            if FinanceBill18=="yes" or FinanceBill18=="y":
                FinanceBill18="https://home.kpmg.com/content/dam/kpmg/ke/pdf/tax/Finance%20Bill%202018%20Analysis.pdf"
                r = requests.get(FinanceBill18, allow_redirects=True)
                FinanceBill18= "Downloading analysis of the Finance Bill 2018 as *Finance Bill 2018-KPMG Analysis*."
                list=textwrap3.wrap(FinanceBill18, width=90)
                for element in list:
                    print(element)
                print("----Download complete----.")
            print("************************************************************")
            print("\n")
            open('Finance Bill 2018-KPMG Analysis.pdf', 'wb').write(r.content)

      ###          
        if actTitle=="Miscellaneous Fees and Levies" or actTitle=="fees and levies" or actTitle=="miscellaneous fees and levies" or actTitle=="fees levies" or actTitle=="levies" or actTitle==" fee and levy" or actTitle=="levy":
         
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=No.%2029%20of%202016"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")

            
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)
                    
            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)
            print("\n")
            #time = soup.find("select", id="cboPIT")
            #print("Amendment History")
            #print(time.text)  
            print(User_Name, ", this is the time stamp of the legislation that matches your keyword: ")
            levy2= "Recent amendments *few*. October 1 2018 see Finance Bill (National Assembly Bills 20) 2018. s.2-'Special Economic Zone'-Insert, First Schedule Part I-Insert, Second Schedule Part A-Insert, Part B-Insert "
            list=textwrap3.wrap(levy2, width=95)
            for element in list:
                print(element)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/M/Miscellaneous%20Fees%20and%20Levies%20Act%20-%20No.%2029%20of%202016/docs/MiscellaneousFeesandLeviesActNo29of2016.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("-----Downloaded Online Version-----.")
            print("\n")
            open('Miscellaneous Fees and Levies .pdf', 'wb').write(r.content)
            print("-"*50)
            FinanceBill18=input("Would you like to view an analysis of the Finance Bill 2018? [y/n] ").lower()
            if FinanceBill18=="yes" or FinanceBill18=="y":
                FinanceBill18="https://home.kpmg.com/content/dam/kpmg/ke/pdf/tax/Finance%20Bill%202018%20Analysis.pdf"
                r = requests.get(FinanceBill18, allow_redirects=True)
                FinanceBill18= "Downloading analysis of the Finance Bill 2018 as *Finance Bill 2018-KPMG Analysis*."
                list=textwrap3.wrap(FinanceBill18, width=90)
                for element in list:
                    print(element)
                print("----Download complete----.")
                open('Finance Bill 2018-KPMG Analysis.pdf', 'wb').write(r.content)
            print("-"*50)
            print("\n")

    ###
        if actTitle=="Valuers" or actTitle=="valuers" or actTitle=="valuer":
        
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%20532"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")

            
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)
            print("\n")
            #time = soup.find("select", id="cboPIT")
            #print("Amendment History")
            #print(time.text)  
            print(User_Name, ", this is the time stamp of the legislation that matches your keyword: ")
            valuer2= "Last amendments *few*. January 1 2001 see Finance Act 9/2000. Section 8A-Repealed. "
            list=textwrap3.wrap(valuer2, width=90)
            for element in list:
                print(element)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/V/Valuers%20Act%20Cap.%20532%20-%20No.%2016%20of%201984/docs/ValuersAct16of1984.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("-----Downloaded Online Version-----.")
            print("\n")
            open('Valuers .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Veterinary Surgeons and Veterinary Para-Professionals" or actTitle=="veterinary" or actTitle=="veterinary surgeon" or actTitle=="vet":
         
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=No.%2029%20of%202011"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")

            
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)
            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)
            print("\n")
            #time = soup.find("select", id="cboPIT")
            #print("Amendment History")
            #print(time.text)  
            print(User_Name, ", this is the time stamp of the legislation that matches your keyword: ")
            vet2= "Last amendments *few*. November 24 2014 see Statute Law (Miscellaneous Amendments) Act 2014. s. 2-Insert, 23(3)-delete, 44A-Insert. "
            list=textwrap3.wrap(vet2, width=90)
            for element in list:
                print(element)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/V/Veterinary%20Surgeons%20and%20Veterinary%20Para-Professionals%20Act%20Cap.%20366%20-%20No.%2029%20of%202011/docs/VeterinarySurgeonsandVeterinaryProfessionalsAct29of2011.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("-----Downloaded Online Version-----.")
            print("\n")
            open('Veterinary Surgeons and Veterinary Para-professionals .pdf', 'wb').write(r.content)

     ###           
        if actTitle=="Vetting of Judges and Magistrates" or actTitle=="vetting" or actTitle=="Vetting of Judges" or actTitle=="Judge" or actTitle=="Magistrate" or actTitle=="judge" or actTitle=="magistrate":
            
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=No.%202%20of%202011"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")

            
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)
            print("\n")
            #time = soup.find("select", id="cboPIT")
            #print("Amendment History")
            #print(time.text)  
            print(User_Name, ", this is the time stamp of the legislation that matches your keyword: ")
            
            vetting2= "Last amendments *few*. January 10 2014 see Vetting of Judges and Magistrates (Amendment) Act. 43/2013. s. 22-substitute, 23-substitute . "
            list=textwrap3.wrap(vetting2, width=90)
            for element in list:
                print(element)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/V/Vetting%20of%20Judges%20and%20Magistrates%20Cap.%208B%20-%20Act%20No.%202%20of%202011/docs/VettingofJudgesandMagistratesAct2of2011.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("-----Downloaded Online Version-----.")
            print("\n")
            
            open('Vetting of Judges and Magistrates .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Vexatious Proceedings" or actTitle=="vexatious proceedings" or actTitle=="vexatious" or actTitle=="Vexatious":
           
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%2041"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")

            
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)
            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)
            ##time = soup.find("select", id="cboPIT")
            ##print("Amendment History")
            ##print(time.text)
            print("\n")
            print(User_Name, ", this is the time stamp of the legislation that matches your keyword: ")
            vex2= "No amendments recorded. "
            list=textwrap3.wrap(vex2, width=90)
            for element in list:
                print(element)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/V/Vexatious%20Proceedings%20Act%20Cap.%2041%20-%20No.%204%20of%201958/docs/VexatiousProceedingsAct4of1958.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("-----Downloaded Online Version-----.")
            print("\n")
            open('Vexatious Proceedings .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Victim Protection" or actTitle=="Victim Protection" or actTitle=="vp" or actTitle=="victim":
           
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=No.%2017%20of%202014"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")

            
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)
            print("\n")
            #time = soup.find("select", id="cboPIT")
            #print("Amendment History")
            #print(time.text)  
            print(User_Name, ", this is the time stamp of the legislation that matches your keyword: ")
            
            vp2= "Last amendments *few*. April 10 2015 see Revision of Laws Rectification Order 2015. s. 31(3)-substitute."
            list=textwrap3.wrap(vp2, width=90)
            for element in list:
                print(element)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/V/Victim%20Protection%20Act%20No%2017%20of%202014/docs/VictimProtectionAct17of2014.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("-----Downloaded Online Version-----.")
            print("\n")
            open('Victim Protection .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Judicature" or actTitle=="judicature":
           
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%208"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")

            
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)
            print("\n")
            #time = soup.find("select", id="cboPIT")
            #print("Amendment History")
            #print(time.text)  
            print(User_Name, ", this is the time stamp of the legislation that matches your keyword: ")
            
            judicature2= "Recent amendments *few*.  May 4 2017 see Statute Law (Miscellaneous Amendments) Act 11/2017. s. 9-substitute. "
            list=textwrap3.wrap(judicature2, width=90)
            for element in list:
                print(element)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/J/Judicature%20Act%20Cap.%208%20-%20No.%2016%20of%201967/docs/JudicatureAct16of1967.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("-----Downloaded Online Version-----.")
            print("\n")
            
            open('Judicature .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Judicial Service" or actTitle=="judicial service" or actTitle=="Judicial service" or actTitle=="Judicial" or actTitle=="judicial":
        
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=No.%201%20of%202011"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")

            
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)
            print("\n")
            #time = soup.find("select", id="cboPIT")
            #print("Amendment History")
            #print(time.text)  
            print(User_Name, ", this is the time stamp of the legislation that matches your keyword: ")
            
            jS2= "Recent amendments *few*. May 4 2017 see Statute Law (Miscellaneous Amendments) Act 11/2017. s. 5(5)-Insert. "
            list=textwrap3.wrap(jS2, width=90)
            for element in list:
                print(element)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/J/Judicial%20Service%20Act%20Cap.%20185B%20-%20Act%20No.%201%20of%202011/docs/JudicialServiceAct1of2011.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("-----Downloaded Online Version-----.")
            print("\n")
            
            open('Judicial Service .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Judiciary Fund" or actTitle=="judiciary fund" or actTitle=="judiciary" or actTitle=="Judiciary":
            
            actTitle1= ["http://www.kenyalaw.org/lex/actview.xql?actid=No.%2016%20of%202016"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")

            
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)

            ##time = soup.find("select", id="cboPIT")
            ##print("Amendment History")
            ##print(time.text)  
            #print(User_Name, ", this is the time stamp of the legislation that matches your keyword: ")
            print("\n")
            jF2= "No amendments recorded. "
            list=textwrap3.wrap(jF2, width=90)
            for element in list:
                print(element)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/J/Judiciary%20Fund%20Act%20No.%2016%20of%202016/docs/JudiciaryFundActNo16of2016.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("-----Downloaded Online Version-----.")
            print("\n")
            
            open('Judiciary Fund .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Cheques" or actTitle=="cheques":
           
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%2035"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")

            
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)

            ##time = soup.find("select", id="cboPIT")
            ##print("Amendment History")
            ##print(time.text)  
            #print(User_Name, ", this is the time stamp of the legislation that matches your keyword: ")
            print("\n")
            cheq2= "Last amendments *few*. See Statute Miscellaneous (Amendments) Act 2/2002. s. 2-substitute. "
            list=textwrap3.wrap(cheq2, width=85)
            for element in list:
                print(element)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/C/Cheques%20Act%20Cap.%2035%20-%20No.%2041%20of%201968/docs/ChequesAct41of1968.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("-----Downloaded Online Version-----.")
            print("\n")
            
            open('Cheques .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Chiefs" or actTitle=="chiefs" or actTitle=="chief" or actTitle=="Chief":
            
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%20128"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
           
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)

            ##time = soup.find("select", id="cboPIT")
            ##print("Amendment History")
            ##print(time.text)  
            #print(User_Name, ", this is the time stamp of the legislation that matches your keyword: ")
            print("\n")
            chief2= "Recent amendments *few*. See Prevention of Torture Act 12/2017. s. 20-delete. "
            list=textwrap3.wrap(chief2, width=85)
            for element in list:
                print(element)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/C/Chiefs%20Act%20Cap.%20128%20-%20No.%202%20of%201937/docs/ChiefsAct2of1937.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("-----Downloaded Online Version-----.")
            print("\n")
            
            open('Chiefs .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Children" or actTitle=="children" or actTitle=="child":
          
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=No.%208%20of%202001"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
            
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)
            print("\n")
            #time = soup.find("select", id="cboPIT")
            #print("Amendment History")
            #print(time.text)  
            print(User_Name, ", this is the time stamp of the legislation that matches your keyword: ")
            
            children2= "Recent amendments *few*. See Statute Miscellaneous Amendments Act 11/2017. s. 31-substitute, 156-Insert. Consequential repeals *few*. Children and Young Persons Act-repealed, Adoption Act-repealed, Guardianship of Infants Act-repealed. "
            list=textwrap3.wrap(children2, width=85)
            for element in list:
                print(element)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/C/Children%20Act%20Cap.%20141%20-%20No.%208%20of%202001/docs/ChildrenAct8of2001.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("-----Downloaded Online Version-----.")
            print("\n")
            
            open('Children .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Dairy Industry" or actTitle=="dairy industry" or actTitle=="dairy" or actTitle=="Dairy":
          
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%20336"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")

            
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)
            print("\n")
            #time = soup.find("select", id="cboPIT")
            #print("Amendment History")
            #print(time.text)  
            print(User_Name, ", this is the time stamp of the legislation that matches your keyword: ")
            
            dairy2= "Last amendments *few*. See Licensing Laws (Repeals and Amendment) Act 17/2006. s. 19(j)-delete.  "
            list=textwrap3.wrap(dairy2, width=85)
            for element in list:
                print(element)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/D/Dairy%20Industry%20Act%20Cap.%20336%20-%20No.%2034%20of%201958/docs/DairyIndustryAct34of1958.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("-----Downloaded Online Version-----.")
            print("\n")
            open('Dairy Industry .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Debts (Summary Recovery)" or actTitle=="debts" or actTitle=="Debts" or actTitle=="debts summary recovery":
            
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%2042"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")

            
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)

            ##time = soup.find("select", id="cboPIT")
            ##print("Amendment History")
            ##print(time.text)
            print("\n")
            print(User_Name, ", this is the time stamp of the legislation that matches your keyword: ")
            debts2= "Last amendments *few*. See Statute Miscellaneous Amendment Act 10/1969. s. 2A-substitute."
            list=textwrap3.wrap(debts2, width=85)
            for element in list:
                print(element)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/D/Debts%20Summary%20Recovery%20Act%20Cap.%2042%20-%20No.%205%20of%201913/docs/DebtsSummaryRecoveryAct5of1913.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("-----Downloaded Online Version-----.")
            print("\n")
            
            open('Debts (Summary Recovery) .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Deeds of Arrangement" or actTitle=="deeds of arrangement" or actTitle=="Deeds" or actTitle=="deeds":
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%2054"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
           
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)

            ##time = soup.find("select", id="cboPIT")
            ##print("Amendment History")
            ##print(time.text)
            print("\n")
            print(User_Name, ", this is the time stamp of the legislation that matches your keyword: ")
            deeds2= "Last amendments *few*. See Kenya Independence Order in Council Act Legal Notice 718/1963 (Kenya (Amendment of Laws)(Written Laws) Order LN 2 of 1964). Throughout-substitute."
            list=textwrap3.wrap(deeds2, width=85)
            for element in list:
                print(element)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/D/Deeds%20of%20Arrangement%20Act%20Cap.%2054%20-%20No.%2038%20of%201930/docs/DeedsofArrangementAct38of1930.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("-----Downloaded Online Version-----.")
            print("\n")
            open('Deeds of Arrangement .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Defamation" or actTitle=="defamation":
          
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%2036"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")

            
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)
            print("\n")
            #time = soup.find("select", id="cboPIT")
            #print("Amendment History")
            #print(time.text)  
            print(User_Name, ", this is the time stamp of the legislation that matches your keyword: ")
            defamation2= "Last amendments *few*. October 23 1992 see Statute Miscellaneous Amendment Act 11/1992. s. 7A-Insert, 16A-Insert, 8(2)-substitute. "
            list=textwrap3.wrap(defamation2, width=85)
            for element in list:
                print(element)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/D/Defamation%20Act%20Cap.%2036%20-%20No.%2010%20of%201970/docs/DefamationAct10of1970.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("-----Downloaded Online Version-----.")
            print("\n")
            
            open('Defamation .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Disposal of uncollected goods" or actTitle=="disposal of uncollected goods" or actTitle=="disposal uncollected goods" or actTitle=="dispose uncollected goods" or actTitle=="uncollected goods":
          
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%2038"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")

            
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)

            ##time = soup.find("select", id="cboPIT")
            ##print("Amendment History")
            ##print(time.text)
            print("\n")
            print(User_Name, ", this is the time stamp of the legislation that matches your keyword: ")
            disposal2= "No amendments recorded. "
            list=textwrap3.wrap(disposal2, width=85)
            for element in list:
                print(element)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/D/Disposal%20of%20Uncollected%20Good%20Act%20Cap.%2038%20-%20No.%203%20of%201987/docs/DisposalofUncollectedGoodAct3of1987.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("-----Downloaded Online Version-----.")
            print("\n")
            open('Disposal of uncollected goods .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Distress for rent" or actTitle=="distress for rent" or actTitle=="Distress for Rent" or actTitle=="rent" or actTitle=="rent distress" or actTitle=="distress rent":
           
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%20293"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")

            
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)

            ##time = soup.find("select", id="cboPIT")
            ##print("Amendment History")
            ##print(time.text)
            print("\n")
            print(User_Name, ", this is the time stamp of the legislation that matches your keyword: ")
            dist2= "Last amendments *few*. See Licensing Laws (Repeals and Amendment) Act 17/2006. s. 18-repealed, Third Schedule-repealed. "
            list=textwrap3.wrap(dist2, width=85)
            for element in list:
                print(element)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/D/Distress%20for%20Rent%20Act%20Cap.%20293%20-%20No.%201%20of%201937/docs/DistressforRentAct1of1937.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("-----Downloaded Online Version-----.")
            print("\n")
            
            open('Distress for Rent .pdf', 'wb').write(r.content)

    ###
        if actTitle=="District and Provinces" or actTitle=="district and provinces" or actTitle=="District" or actTitle=="Provinces":
          
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=No.5%20of%201992"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")

            
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)

            ##time = soup.find("select", id="cboPIT")
            ##print("Amendment History")
            ##print(time.text)
            print("\n")
            print(User_Name, ", this is the time stamp of the legislation that matches your keyword: ")
            prov2= " Last amendments *few*. October 23 1992 see Statute law (MiscellaneoUs Amendments) Act 11/1992. Second Schedule-substitute. "
            list=textwrap3.wrap(prov2, width=85)
            for element in list:
                print(element)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/D/Districts%20and%20Provinces%20Act%20Cap.%20105A%20-%20No.%205%20of%201992/docs/DistrictsAndProvincesAct105A-5of1992.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("-----Downloaded Online Version-----.")
            print("\n")
            
            open('District and Provinces .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Unclaimed Financial Assets" or actTitle=="ufa" or actTitle=="UFA" or actTitle=="unclaimed financial assets" or actTitle=="unclaimed assets":
            
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=No.%2040%20of%202011"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
          
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)
            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)

            ##time = soup.find("select", id="cboPIT")
            ##print("Amendment History")
            ##print(time.text)
            print("\n")
            print(User_Name, ", this is the time stamp of the legislation that matches your keyword: ")
            prov2= "No amendments recorded. "
            list=textwrap3.wrap(prov2, width=85)
            for element in list:
                print(element)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/U/Unclaimed%20Financial%20Assets%20Act%20-%20No.%2040%20of%202011/docs/UnclaimedFinanancialAssetAct40of2011.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("I downloaded a copy just in case.")
            print("\n")
            
            open('Unclaimed Financial Assets .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Universities" or actTitle=="universities" or actTitle=="higher education" or actTitle=="Higher Education" or actTitle=="university" or actTitle=="campus":
           
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=No.%2042%20of%202012"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
          
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)
            print("\n")
            #time = soup.find("select", id="cboPIT")
            #print("Amendment History")
            #print(time.text)
            print("\n")
            print(User_Name, ", this is the time stamp of the legislation that matches your keyword: ")
            prov2= "Last recorded amendments: January 13 2017 see Universities (Amendment) Act 2016. "
            list=textwrap3.wrap(prov2, width=85)
            for element in list:
                print(element)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/U/Universities%20Act%20No.%2042%20of%202012/docs/UniversitiesAct42of2012.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("-----Downloaded Online Version-----.")
            print("\n")
            open('Universities .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Uplands Bacon Factory" or actTitle=="uplands bacon factory" or actTitle=="Uplands Bacon" or actTitle=="uplands bacon" or actTitle=="uplands" or actTitle=="bacon" or actTitle=="Uplands bacon" or actTitle=="Uplands" or actTitle=="Bacon":
            
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%20362"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
          
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)
            print("\n")
            #time = soup.find("select", id="cboPIT")
            #print("Amendment History")
            #print(time.text)
            print("\n")
            print(User_Name, ", this is the time stamp of the legislation that matches your keyword: ")
            prov2= "Last amendments *few*. See Companies and Insolvency Legislation (Consequential Amendments) Act. s. 3(3)-delete. "
            list=textwrap3.wrap(prov2, width=85)
            for element in list:
                print(element)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/U/Uplands%20Bacon%20Factory%20Act%20Cap.%20362%20-%20No.%2020%20of%201945/docs/UplandsBaconFactoryAct%2020of1945.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("-----Downloaded Online Version-----.")
            print("\n")
            open('Uplands Bacon Factory .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Urban Areas and Cities" or actTitle=="urban areas and cities" or actTitle=="Urban Areas" or actTitle=="urban areas" or actTitle=="urban cities" or actTitle=="urban area" or actTitle=="urban":
          
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=No.%2013%20of%202011"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
        
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)
            print("\n")
            #time = soup.find("select", id="cboPIT")
            #print("Amendment History")
            #print(time.text)
            print("\n")
            print(User_Name, ", this is the time stamp of the legislation that matches your keyword: ")
            prov2= "Recent amedments *few*. May 10 2016 see Statute Law (Miscellaneous Amendments) Act 7 of 2016. s. 5-substitute, 9-substitute Insert, 10-substitute, 61-Insert. "
            list=textwrap3.wrap(prov2, width=85)
            for element in list:
                print(element)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/U/Urban%20Areas%20and%20Cities%20Act%20Cap.%20275%20-%20No.%2013%20of%202011/docs/UrbanAreasAndCitiesAct13of2011.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("-----Downloaded Online Version-----.")
            print("\n")
            open('Urban Areas and Cities .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Use of poisonous substances" or actTitle=="poison" or actTitle=="poisonous substances" or actTitle=="Poisonous substances" or actTitle=="Poisonous Substances" or actTitle=="poisonous" :
            
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%20247"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
        
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)
            print("\n")
            #time = soup.find("select", id="cboPIT")
            #print("Amendment History")
            #print(time.text)
            print("\n")
            print(User_Name, ", this is the time stamp of the legislation that matches your keyword: ")
            prov2= "Last amendments *few*. See Licensing Laws (Repeals and Amendment) Act 7/2006. s. 3(1)-delete. "
            list=textwrap3.wrap(prov2, width=85)
            for element in list:
                print(element)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/U/Use%20of%20Poisonous%20Substances%20Act%20Cap.%20247%20-%20No.%2023%20of%201957/docs/UseOfPoisonousSubstancesAct23of1957.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("-----Downloaded Online Version-----.")
            print("\n")
            open('Use of poisonous substances .pdf', 'wb').write(r.content)

    ###
        if actTitle=="WAKF Commissioners" or actTitle=="wakf commissioners" or actTitle=="WAKF" or actTitle=="wakf":
           
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%20109"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")

            
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)

            ##time = soup.find("select", id="cboPIT")
            ##print("Amendment History")
            ##print(time.text)
            print("\n")
            print(User_Name, ", this is the time stamp of the legislation that matches your keyword: ")
            prov2= "Last amendments *few*. June 11 1971 see Statute Law (Miscellaneous Amendments) Act 14/1971. s. 6(1)-substitute, 6(2)-substitute, 8(1)-substitute, 9-Insert, 23-substitute, 25(2)-substitute. "
            list=textwrap3.wrap(prov2, width=85)
            for element in list:
                print(element)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/W/Wakf%20Commissioners%20Act%20Cap.%20109%20-%20No.%2030%20of%201951/docs/WakfCommissionersAct30of1951.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("-----Downloaded Online Version-----.")
            print("\n")
            open('Wakf Commissioners .pdf', 'wb').write(r.content)

            wakfBill= "http://www.statelaw.go.ke/wakf-bill-2018-submitted-to-ag/"
            r = requests.get(wakfBill, allow_redirects=True)
            print("Update.")
            print("WAKF Bill 2018 submitted to ag")
            print("check folder for details")
            open('Wakf Bill Statement AG.html', 'wb').write(r.content)

    ###
        if actTitle=="Water" or actTitle=="water":
           
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=No.%2043%20of%202016"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
           
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)
            print("\n")
            #time = soup.find("select", id="cboPIT")
            #print("Amendment History")
            #print(time.text)
            print("\n")
            print(User_Name, ", this is the time stamp of the legislation that matches your keyword: ")
            prov2= "Recent amendments *few*. May 4 2017 see Statute Law (Miscellaneous Amendments) Act 11/2017. s. 37(4)-Insert, 74-Insert and 85 (3)(d)-Insert."
            list=textwrap3.wrap(prov2, width=85)
            for element in list:
                print(element)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/W/Water%20Act%20-%20No.%2043%20of%202016/docs/WaterAct43of2016.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("-----Downloaded Online Version-----.")
            print("\n")
            
            open('Water .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Weights and Measures" or actTitle=="Weights & Measures" or actTitle=="weights & measures" or actTitle=="Weights Measures" or actTitle=="weights measures" or actTitle=="Weights" or actTitle=="Measures" or actTitle=="weights" or actTitle=="measures":
            
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%20513"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
          
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)
            print("\n")
            #time = soup.find("select", id="cboPIT")
            #print("Amendment History")
            #print(time.text)
            print("\n")
            print(User_Name, ", this is the time stamp of the legislation that matches your keyword: ")
            prov2= "Last amendments *many*. June 7 2002 see Statute Law (Miscellaneous Amendments) Act 2/2002."
            list=textwrap3.wrap(prov2, width=85)
            for element in list:
                print(element)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/W/Weights%20and%20Measures%20Act%20Cap.%20513%20-%20No.%2018%20of%201987/docs/WeightsandMeasuresAct18of1987.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("-----Downloaded Online Version-----.")
            print("\n")
            open('Weights & Measures .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Widows and Childrens Pensions" or actTitle=="widows and childrens pensions" or actTitle=="Widows Childrens Pensions" or actTitle=="widows childrens pensions" or actTitle=="widows" or actTitle=="widows and children":
           
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%20195"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
          
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)

            ##time = soup.find("select", id="cboPIT")
            ##print("Amendment History")
            ##print(time.text)
            print("\n")
            print(User_Name, ", this is the time stamp of the legislation that matches your keyword: ")
            prov2= "Last amendments *few*. See Public Service Superannuation Scheme Act Chapter 190A. s. 3(1)(a)-Insert."
            list=textwrap3.wrap(prov2, width=85)
            for element in list:
                print(element)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/W/Widows%20and%20Childrens%20Pensions%20Act%20Cap.%20195%20-%20No.%2031%20of%201965/docs/WidowsandChildrensPensionsAct31of1965.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("-----Downloaded Online Version-----.")
            print("\n")
            
            open('Widows and Childrens Pensions .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Widows and Orphans Pensions" or actTitle=="widows and orphans pensions" or actTitle=="Widows Orphans Pensions" or actTitle=="widows orphans pensions" or actTitle=="widows" or actTitle=="widows and orphans":
            
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%20192"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
            
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)

            ##time = soup.find("select", id="cboPIT")
            ##print("Amendment History")
            ##print(time.text)
            print("\n")
            print(User_Name, ", this is the time stamp of the legislation that matches your keyword: ")
            prov2= "Last amendments *few*. July 12 1966 see Statute Laws (Miscellaneous Amendment) Act 21 of 1966. s. 37(4)-delete."
            list=textwrap3.wrap(prov2, width=85)
            for element in list:
                print(element)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/W/Widows%20and%20Orphans%20Pensions%20Act%20Cap.%20192%20-%20No.%2022%20of%201921/docs/WidowsandOrphansPensionsAct22of1921.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("-----Downloaded Online Version-----.")
            print("\n")
            
            open('Widows and Orphans Pensions .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Wildlife Conservation and Management" or actTitle=="wildlife conservation and management" or actTitle=="Wildlife Conservation Management" or actTitle=="wildlife conservation management" or actTitle=="wildlife conservation" or actTitle=="wildlife":
            
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=No.%2047%20of%202013"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")

            
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)

            ##time = soup.find("select", id="cboPIT")
            ##print("Amendment History")
            ##print(time.text)
            print("\n")
            print(User_Name, ", this is the time stamp of the legislation that matches your keyword: ")
            prov2= "No amendments recorded."
            list=textwrap3.wrap(prov2, width=85)
            for element in list:
                print(element)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/W/Wildlife%20Conservation%20and%20Management%20Act%20-%20No.%2047%20of%202013/docs/WildlifeConservationandManagementAct(ActNo.47of2013).pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("I downloaded a copy just in case. ..")
            print("\n")
            
            open('Wildlife Conservation and Management .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Witchcraft" or actTitle=="witchcraft" or actTitle=="witch craft" or actTitle=="Witch Craft" or actTitle=="Witch craft":
            
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%2067"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")

            
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)

            ##time = soup.find("select", id="cboPIT")
            ##print("Amendment History")
            ##print(time.text)
            print("\n")
            print(User_Name, ", this is the time stamp of the legislation that matches your keyword: ")
            prov2= "Last amendment *unknown*. April 24 1964 see Kenya (Amendment of Laws) (Miscellaneous Amendments) (No. 2) Order of 1964."
            list=textwrap3.wrap(prov2, width=85)
            for element in list:
                print(element)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/W/Witchcraft%20Act%20Cap.%2067/docs/WitchcraftAct23of1925.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("I downloaded a copy just in case. ..")
            print("\n")
            open('Witchcraft .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Witness Protection" or actTitle=="witness protection" or actTitle=="Witness" or actTitle=="witness":
            
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=No.%2016%20of%202006"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")

            
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)
            print("\n")
            #time = soup.find("select", id="cboPIT")
            #print("Amendment History")
            #print(time.text)
            print("\n")
            print(User_Name, ", this is the time stamp of the legislation that matches your keyword: ")
            prov2= "Last recorded amendments: January 13 2017 see Witness Protection (Amendment) Act 45 of 2016."
            list=textwrap3.wrap(prov2, width=85)
            for element in list:
                print(element)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/W/Witness%20Protection%20Act%20Cap.%2079%20-%20No.%2016%20of%202006/docs/WitnessProtectionAct16of2006.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("I downloaded a copy just in case. ..")
            print("\n")
            
            open('Witness Protection .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Witness Summonses (Reciprocal Enforcement)" or actTitle=="witness summonses (reciprocal enforcement)" or actTitle=="witness summonses" or actTitle=="Witness Summonses":
          
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%2078"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")

            
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)

            ##time = soup.find("select", id="cboPIT")
            ##print("Amendment History")
            ##print(time.text)
            print("\n")
            print(User_Name, ", this is the time stamp of the legislation that matches your keyword: ")
            prov2= "No amendments recorded. "
            list=textwrap3.wrap(prov2, width=85)
            for element in list:
                print(element)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/W/Witness%20Summonses%20Reciprocal%20Enforcement%20Act%20Cap.%2078%20-%20No.%2062%20of%201968/docs/WitnessSummonsesReciprocal%20EnforcementAct62of1968.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("I downloaded a copy just in case. ..")
            print("\n")
            open('Witness Summonses .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Work Injury Benefits" or actTitle=="works injury benefits" or actTitle=="injury benefits" or actTitle=="wiba" or actTitle=="WIBA":
            
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=No.%2013%20of%202007"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")

            
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)

            ##time = soup.find("select", id="cboPIT")
            ##print("Amendment History")
            ##print(time.text)
            print("\n")
            print(User_Name, ", this is the time stamp of the legislation that matches your keyword: ")
            prov2= "No amendments recorded. Possible Nexus - Factories and Other Places of Work Act Chapter 514. "
            list=textwrap3.wrap(prov2, width=85)
            for element in list:
                print(element)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/W/Work%20Injury%20Benefits%20Act%20Cap.%20236%20-%20No.%2013%20of%202007/docs/WorkInjuryBenefitsAct13of2007.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("I downloaded a copy just in case. ..")
            print("\n")
            open('Work Injury Benefits .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Parliamentary Pensions" or actTitle=="parliamentary pensions" or actTitle=="Parliament pensions" or actTitle=="parliament pensions" or actTitle=="parliament" or actTitle=="Parliament" or actTitle=="Parliamentary" or actTitle=="parliamentary":
           
            actTitle1= ["http://kenyalaw.org/lex//actview.xql?actid=CAP.%20196"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
           
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)
            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)
            print("\n")
            #time = soup.find("select", id="cboPIT")
            #print("Amendment History")
            #print(time.text)
            print("\n")
            print(User_Name, ", this is the time stamp of the legislation that matches your keyword: ")
            prov2= "Last amendment *few*. January 1 2001 see Statutes Law (Miscellaneous Amendments) Act 2 of 2002.  s. 3(1)-delete, 5-Substitute, 7(1)-Insert, Substitute, 8(1)-Substitute, 9a-Substitute. "
            list=textwrap3.wrap(prov2, width=85)
            for element in list:
                print(element)

            Download= "http://kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/P/Parliamentary%20Pensions%20Act%20Cap.%20196%20-%20No.%204%20of%201983/docs/ParliamentaryPensionsAct4of1983.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("I downloaded a copy just in case. ..")
            print("\n")
            
            open('Parliamentary Pensions .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Parliamentary Powers and Privileges" or actTitle=="parliamentary powers and privileges" or actTitle=="parliamentary powers privileges" or actTitle=="parliamentary powers"  or actTitle=="parliamentary privileges" or actTitle=="Parliament" or actTitle=="Parliamentary" or actTitle=="parliamentary":
            
            actTitle1= ["http://kenyalaw.org/lex//actview.xql?actid=No.%2029%20of%202017"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
            
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)
            print("\n")
            #time = soup.find("select", id="cboPIT")
            #print("Amendment History")
            #print(time.text)
            print("\n")
            print(User_Name, ", this is the time stamp of the legislation that matches your keyword: ")
            prov2= "No amendments recorded. "
            list=textwrap3.wrap(prov2, width=85)
            for element in list:
                print(element)

            Download= "http://kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/P/Parliamentary%20Powers%20and%20Privileges%20Act%20-%20No.%2029%20of%202017/docs/ParliamentaryPowersandPrivelegesAct29of2017.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("I downloaded a copy just in case. ..")
            print("\n")
            
            open('Parliamentary Powers and Privileges .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Parliamentary Service" or actTitle=="parliamentary service":
           
            actTitle1= ["http://kenyalaw.org/lex//actview.xql?actid=No.%2010%20of%202000"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
            
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)
            print("\n")
            #time = soup.find("select", id="cboPIT")
            #print("Amendment History")
            #print(time.text)
            print("\n")
            print(User_Name, ", this is the time stamp of the legislation that matches your keyword: ")
            prov2= "Recent amendments *few*.  See Public Procurement Act 3 of 2005 (now repealed) Fourth schedule s. 22-delete. "
            list=textwrap3.wrap(prov2, width=85)
            for element in list:
                print(element)

            Download= "http://kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/P/Parliamentary%20Service%20Act%20Cap.%20185A%20-%20No.%2010%20of%202000/docs/ParliamentaryServiceAct10of2000.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("I downloaded a copy just in case. ..")
            print("\n")
            open('Parliamentary Service .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Partnerships" or actTitle=="partnerships" or actTitle=="Partnership" or actTitle=="partnership":
            
            actTitle1= ["http://kenyalaw.org/lex//actview.xql?actid=No.%2016%20of%202012"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
            
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)

            ##time = soup.find("select", id="cboPIT")
            ##print("Amendment History")
            ##print(time.text)
            print("\n")
            print(User_Name, ", this is the time stamp of the legislation that matches your keyword: ")
            prov2= "Last amendment *few*. Parternships Act Chapter 29-Repealed in section 77. No amendments recorded in current. "
            list=textwrap3.wrap(prov2, width=85)
            for element in list:
                print(element)

            Download= "http://kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/P/Partnerships%20Act%20No.%2016%20of%202012/docs/PartnershipsAct16of2012.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("I downloaded a copy just in case.")
            print("\n")       
            open('Partnerships .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Penalcode" or actTitle=="penalcode" or actTitle=="Penalcode" or actTitle=="Penal Code" or actTitle=="penal code" or actTitle=="penal" or actTitle=="Penal" or actTitle=="Code" or actTitle=="code":
           
            actTitle1= ["http://kenyalaw.org/lex//actview.xql?actid=CAP.%2063"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
           
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)
            print("\n")
            #time = soup.find("select", id="cboPIT")
            #print("Amendment History")
            #print(time.text)
            print("\n")
            print(User_Name, ", this is the time stamp of the legislation that matches your keyword: ")
            prov2= "Last amendments *few*. December 22 2014 see Security Laws (Amendment) Act 19/2014. s. 66A-Insert, 128A-Insert, 251A-Insert. See also Penal Code (Amendment) Bill 13 of 2016 (Kenya Gazette Supplement No. 54 of 2016). Bill Tracker status= Fail."
            list=textwrap3.wrap(prov2, width=85)
            for element in list:
                print(element)

            Download="http://kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/P/Penal%20Code%20Cap.%2063%20-%20No.%2010%20of%201930/docs/PenalCode81of1948.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("I downloaded a copy just in case. ..")
            print("\n")
            open('Penal Code .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Pensions (Increase)" or actTitle=="pensions increase" or actTitle=="increase pensions" or actTitle=="increase pension" or actTitle=="pension increase" or actTitle=="increased pensions" or actTitle=="increased pension":
            
            actTitle1= ["http://kenyalaw.org/lex//actview.xql?actid=CAP.%20190"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
            
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)
            print("\n")
            #time = soup.find("select", id="cboPIT")
            #print("Amendment History")
            #print(time.text)
            print("\n")
            print(User_Name, ", this is the time stamp of the legislation that matches your keyword: ")
            prov2= "Last amendments *few*. See Finance Act 9/2007. s. 3-substitute, 12-substitute, Second schedule-repealed. "
            list=textwrap3.wrap(prov2, width=85)
            for element in list:
                print(element)

            Download="http://kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/P/Pensions%20Increase%20Act%20Cap.%20190%20-%20No.%2010%20of%201957/docs/PensionsIncreaseAct10of1957.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("I downloaded a copy just in case. ..")
            print("\n")
            
            open('Pensions (Increase) .pdf', 'wb').write(r.content)

    ###
        if actTitle=="PPP" or actTitle=="PUBLIC PRIVATE PARTNERSHIPS" or actTitle=="ppp":
          
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=No.%2015%20of%202013"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
            
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)
            print("\n")
            #time = soup.find("select", id="cboPIT")
            #print("Amendment History")
            #print(time.text)
            print("\n")
            print(User_Name, ", this is the time stamp of the legislation that matches your keyword: ")
            prov2= "Last amendments *few*. October 1 2015 see Finance Act 14 of 2015. s. 4(1)(b)-delete, 4 (i, j, k)-insert, 43(gg)-Insert. "
            list=textwrap3.wrap(prov2, width=85)
            for element in list:
                print(element)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/P/Public%20Private%20Partnerships%20Act%20No.%2015%20of%202013/docs/PublicPrivatePartnershipsActNo.15of2013.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("I downloaded a copy just in case. ..")
            print("\n")
            open('Public Private Partnership .pdf', 'wb').write(r.content)

            Download1=input ("Do you need the Public Private Partnerships status report of 2018? [y/n] ").lower()
            if Download1=="yes" or Download1=="y":
                Download1= "https://pppunit.go.ke/wp-content/uploads/2018/06/Kenya-PPP-Pipeline-Status-Report-June-2018.pdf"
                r = requests.get(Download1, allow_redirects=True)
                Download1= "Downloading Public Private Partnerships status report of 2018 as *PPP status report-June 2018*."
                list=textwrap3.wrap(Download1, width=90)
                for element in list:
                    print(element)
                print("----Download complete----.")
                open('PPP status report-June 2018.pdf', 'wb').write(r.content)

    ###
        if actTitle=="Movable Property Security Rights" or actTitle=="movable property" or actTitle=="movable property security rights" or actTitle=="MPSR":
            
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=No.%2013%20of%202017"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
            
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)
            print("\n")
            #time = soup.find("select", id="cboPIT")
            #print("Amendment History")
            #print(time.text)
            print("\n")
            print(User_Name, ", this is the time stamp of the legislation that matches your keyword: ")
            prov2= "No amendments recorded. Consequenstial amendments and repeals *few*. Chattels Transfer Act-Repealed, Agricultural Finance Corporation Act s.2- Insert, Stamp Duty s. 31-Delete s.38-Delete s.117-Insert, Hire Purchase Act s.2-delete s.4 5 35(2)-Delete, Pawnbrokers Act-Repealed, Business Registration Service s. 4-substitute. "
            list=textwrap3.wrap(prov2, width=85)
            for element in list:
                print(element)

            Download= "http://www.statelaw.go.ke/wp-content/uploads/2016/07/13%E2%80%94Movable-Property-Security-Rights-Act-2017-Full.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("-----Downloaded Online Version-----.")
            print("\n")
            open('Movable Property Security Rights .pdf', 'wb').write(r.content)

            Download1=input ("Do you need the Movable Property Security Rights (General) Regulations? [y/n] ").lower()
            if Download1=="yes" or Download1=="y":
                Download1= "http://kenyalaw.org/kl/fileadmin/pdfdownloads/LegalNotices/2017/86-MovablePropertySecurityRights_General_Regulations_2017.pdf"
                r = requests.get(Download1, allow_redirects=True)
                Download1= "Downloading Movable Property Security Rights (General) Regulations as *MPSR regulations*."
                list=textwrap3.wrap(Download1, width=90)
                for element in list:
                    print(element)
                print("----Download complete----.")
                open('MPSR regulations.pdf', 'wb').write(r.content)

    ###
        if actTitle=="Agricultural Finance Corporation" or actTitle=="afc":
            
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%20323"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
            
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)
            print("\n")
            #time = soup.find("select", id="cboPIT")
            #print("Amendment History")
            #print(time.text)
            print("\n")
            print(User_Name, ", this is the time stamp of the legislation that matches your keyword: ")
            prov2= "Recent amendments *few*. See schedule to Movable Property Security Rights Act 13/2017. s.2-Insert. "
            list=textwrap3.wrap(prov2, width=85)
            for element in list:
                print(element)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/A/Agricultural%20Finance%20Corporation%20Act%20Cap.%20323%20-%20No.%201%20of%201969/docs/AgriculturalFinanceCorporationAct1of1969.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("-----Downloaded Online Version-----.")
            print("\n")
            open('Agricultural Finance Corporation .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Pensions" or actTitle=="pensions" or actTitle=="pension" or actTitle=="Pension":
             
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%20189"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")

            
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)
            print("\n")
            #time = soup.find("select", id="cboPIT")
            #print("Amendment History")
            #print(time.text)
            print("\n")
            print(User_Name, ", this is the time stamp of the legislation that matches your keyword: ")
            prov2= "Recent amendments *few*. May 21 2018 see Statute Law (Miscellaneous Amendments) Act. s. 2-Insert, 5(3)-substitute. "
            list=textwrap3.wrap(prov2, width=85)
            for element in list:
                print(element)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/P/Pensions%20Act%20Cap.%20189%20-%20No.%2031%20of%201950/docs/PensionsAct31of1950.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("-----Downloaded Online Version-----.")
            print("\n")
            open('Pensions .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Pharamcy and Poisons" or actTitle=="pharmacy and poisons" or actTitle=="Pharmacy" or actTitle=="pharmacy" or actTitle=="poisons" or actTitle=="Poisons":
           
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%20244"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
           
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)
            print("\n")
            #time = soup.find("select", id="cboPIT")
            #print("Amendment History")
            #print(time.text)
            print("\n")
            print(User_Name, ", this is the time stamp of the legislation that matches your keyword: ")
            prov2= "Recent amendments *few*. May 21 2018 see Statute Law (Miscellaneous Amendments) Act. s. 5(1)-substitute. "
            list=textwrap3.wrap(prov2, width=85)
            for element in list:
                print(element)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/P/Pharmacy%20and%20Poisons%20Act%20Cap.%20244%20-%20No.%2017%20of%201956/docs/PharmacyandPoisonsAct17of1956.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("-----Downloaded Online Version-----.")
            print("\n")
            open('Pharmacy and Poisons .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Environmental Management and Co-ordination" or actTitle=="environmental management and co-ordination" or actTitle=="emca" or actTitle=="EMCA" or actTitle=="Environmental Management" or actTitle=="environmental management" or actTitle=="environmental":
            
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=No.%208%20of%201999"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
           
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)
            print("\n")
            #time = soup.find("select", id="cboPIT")
            #print("Amendment History")
            #print(time.text)
            print("\n")
            print(User_Name, ", this is the time stamp of the legislation that matches your keyword: ")
            prov2= "Recent amendments *few*. May 21 2018 see Statute Law (Miscellaneous Amendments) Act 4/2018. s. 125(1)-substitute, 125(5)-substitute, 129-substitute. "
            list=textwrap3.wrap(prov2, width=85)
            for element in list:
                print(element)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/E/Environmental%20Management%20and%20Co-ordination%20Act%20Cap.%20387%20-%20No.%208%20of%201999/docs/EnvironmentalManagementandCo-ordinationAct8of1999.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("-----Downloaded Online Version-----.")
            print("\n")
            open('Environmental Management and Co-ordination .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Salaries and Remuneration Commission" or actTitle=="salaries and remuneration commission" or actTitle=="Salaries and Remuneration" or actTitle=="salaries":
            
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=No.%2010%20of%202011"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
            
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)
            print("\n")
            #time = soup.find("select", id="cboPIT")
            #print("Amendment History")
            #print(time.text)
            print("\n")
            print(User_Name, ", this is the time stamp of the legislation that matches your keyword: ")
            prov2= "Recent amendments *few*. May 21 2018 see Statute Law (Miscellaneous Amendments) Act 4/2018. s. 4(3)-substitute, 7(3,4,5,6,7,8,9)-delete, 7(10)-substitute, 7(11)-substitute. "
            list=textwrap3.wrap(prov2, width=85)
            for element in list:
                print(element)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/S/Salaries%20and%20Remuneration%20Commission%20Act%20Cap.%205F%20-%20No.%2010%20of%202011/docs/SalariesandRemunerationCommissionAct10of2011.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("-----Downloaded Online Version-----.")
            print("\n")
            open('Salaries and Remuneration Commission .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Statutory Instruments" or actTitle=="statutory instruments" or actTitle=="statutory":
           
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=No.%2023%20of%202013"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
           
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)
            print("\n")
            #time = soup.find("select", id="cboPIT")
            #print("Amendment History")
            #print(time.text)
            print("\n")
            print(User_Name, ", this is the time stamp of the legislation that matches your keyword: ")
            prov2= "Recent amendments *few*. May 21 2018 see Statute Law (Miscellaneous Amendments) Act 4/2018. s. 2-substitute, 5(A)-substitute, 11A-Insert, 11(2)-substitute. 11(4)-Insert, 15(2)-Insert. "
            list=textwrap3.wrap(prov2, width=85)
            for element in list:
                print(element)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/S/Statutory%20Instruments%20Act%20No.%2023%20of%202013/docs/StatutoryInstrumentsActNo.23of2013.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("-----Downloaded Online Version-----.")
            print("\n")
            open('Statutory Instruments .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Occupational Therapists(Training, Registration and Licensing)" or actTitle=="therapist":
           
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=No.%2031%20of%202017"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
            
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)
            print("\n")
            #time = soup.find("select", id="cboPIT")
            #print("Amendment History")
            #print(time.text)
            print("\n")
            print(User_Name, ", this is the time stamp of the legislation that matches your keyword: ")
            prov2= "Recent amendments *few*. May 21 2018 see Statute Law (Miscellaneous Amendments) Act 4/2018. s. 4(1)-substitute. "
            list=textwrap3.wrap(prov2, width=85)
            for element in list:
                print(element)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/S/Statutory%20Instruments%20Act%20No.%2023%20of%202013/docs/StatutoryInstrumentsActNo.23of2013.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("-----Downloaded Online Version-----.")
            print("\n")
            open('Statutory Instruments .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Public Trustee" or actTitle=="public trustee" or actTitle=="public trustees" or actTitle=="Public Trustees":
            
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%20168"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
           
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)
            print("\n")
            #time = soup.find("select", id="cboPIT")
            #print("Amendment History")
            #print(time.text)
            print("\n")
            print(User_Name, ", this is the time stamp of the legislation that matches your keyword: ")
            prov2= "Last recorded amendments: June 8 2018 see Public Trustee (Amendment)  Act 6/2018. "
            list=textwrap3.wrap(prov2, width=85)
            for element in list:
                print(element)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/P/Public%20Trustee%20Act%20Cap.%20168%20-%20No.%2015%20of%201951/docs/PublicTrusteeAct.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("-----Downloaded Online Version-----.")
            print("\n")
            open('Public Trustee .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Perpetuities and Accumulations" or actTitle=="perpetuities and accumulations" or actTitle=="perpetuities" or actTitle=="accumulations":
           
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%20161"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
            
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)

            ##time = soup.find("select", id="cboPIT")
            ##print("Amendment History")
            ##print(time.text)
            print("\n")
            print(User_Name, ", this is the time stamp of the legislation that matches your keyword: ")
            prov2= "No amendments recorded. "
            list=textwrap3.wrap(prov2, width=85)
            for element in list:
                print(element)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/P/Perpetuities%20and%20Accumulations%20Act%20Cap.%20161%20-%20No.%206%20of%201984/docs/PerpetuitiesandAccumulationsAct6of1984.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("-----Downloaded Online Version-----.")
            print("\n")
            open('Perpetuities and Accumulations .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Persons deprived liberty" or actTitle=="persons deprived liberty" or actTitle=="deprive" or actTitle=="liberty":
            
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=No.%2023%20of%202014"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
           
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)

            ##time = soup.find("select", id="cboPIT")
            ##print("Amendment History")
            ##print(time.text)
            print("\n")
            print(User_Name, ", this is the time stamp of the legislation that matches your keyword: ")
            prov2= "No amendments recorded. "
            list=textwrap3.wrap(prov2, width=85)
            for element in list:
                print(element)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/P/Persons%20Deprived%20of%20Liberty%20Act%20No.%2023%20of%202014/docs/PersonsDeprivedofLibertyAct2014.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("-----Downloaded Online Version-----.")
            print("\n")
            open('Persons Deprived of Liberty .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Persons with disabilities" or actTitle=="persons with disabilities" or actTitle=="disabilities" or actTitle=="disability":
            
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=No.%2014%20of%202003"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
            
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)

            ##time = soup.find("select", id="cboPIT")
            ##print("Amendment History")
            ##print(time.text)
            print("\n")
            print(User_Name, ", this is the time stamp of the legislation that matches your keyword: ")
            prov2= "Recent amendments *few*. May 4 2017 see Statute Law (Miscellaneous Amendments) Act 11/2017. s.4-substitute, 5-substitute. "
            list=textwrap3.wrap(prov2, width=85)
            for element in list:
                print(element)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/P/Persons%20with%20Disabilities%20Act%20No.%2014%20of%202003/docs/PersonswithDisabilitiesAct14of2003.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("-----Downloaded Online Version-----.")
            print("\n")
            open('Persons with disabilities .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Computer Misuse and Cybercrime" or actTitle=="cybercrime" or actTitle=="computer" or actTitle=="hacking" or actTitle=="stalking" or actTitle=="cyber" or actTitle=="technology":
            
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=No.%205%20of%202018"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
            
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)
            print("\n")
            #time = soup.find("select", id="cboPIT")
            #print("Amendment History")
            #print(time.text)
            print("\n")
            print(User_Name, ", this is the time stamp of the legislation that matches your keyword: ")
            print("No amendments recorded. ")
            
            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/C/Computer%20Misuse%20and%20Cybercrimes%20Act%20-%20No.%2015%20of%202018/docs/ComputerMisuseandCybercrimesAct5of2018.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("-----Downloaded Online Version-----.")
            print("\n")
            open('Computer Misuse and Cybercrime .pdf', 'wb').write(r.content)                                 
                    
    ###
        if actTitle=="Civil Procedure" or actTitle=="civil procedure" or actTitle=="civil procedure rules" or actTitle=="civil":
          
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%2021"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
           
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            quote_page = "http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%2021"
            page = urllib.request.urlopen(quote_page)
            soup = BeautifulSoup(page, "lxml")
            Amending_Acts=soup.find("div", attrs={"id": "longTitle_1_V5_t5"})
            print(Amending_Acts.text)

            time = soup.find("select", id="cboPIT")
            print("Amendment history")
            print(time.text)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/C/Civil%20Procedure%20Act%20Cap.%2021%20-%20No.%203%20of%201924/docs/CiviProcedureAct3of1924.pdf"
            r = requests.get(Download, allow_redirects=True)
            toaster.show_toast("Downloaded Online Version ",
                                        " ",
                                        icon_path="scrapercon.ico",
                                        duration=5)
            open('Civil Procedure Act .pdf', 'wb').write(r.content)

            Download0=input("Do you need the Civil Procedure Rules 2010? [y/n] ").lower()
            if Download0=="yes" or Download0=="y":
                Download0= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/C/Civil%20Procedure%20Act%20Cap.%2021%20-%20No.%203%20of%201924/subsidiary%20legislation/docs/CiviProcedureAct3of1924_subsidiary.pdf"
                r = requests.get(Download0, allow_redirects=True)
                open('Civil Procedure Rules .pdf', 'wb').write(r.content)
                print(title.string, "rules download complete. ")

    ###
        if actTitle=="Law Society of Kenya" or actTitle=="lsk" or actTitle=="LSK":
          
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=No.%2021%20of%202014"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
            
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            quote_page = "http://www.kenyalaw.org/lex//actview.xql?actid=No.%2021%20of%202014"
            page = urllib.request.urlopen(quote_page)
            soup = BeautifulSoup(page, "lxml")
            Amending_Acts=soup.find("div", attrs={"id": "longTitle_1_V5_t5"})
            print(Amending_Acts.text)

            time = soup.find("select", id="cboPIT")
            print("Amendment history")
            print(time.text)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/L/Law%20Society%20of%20Kenya%20%20Act%20No.%2021%20of%202014/docs/Law_Society_of_Kenya(No21of2014).pdf"
            r = requests.get(Download, allow_redirects=True)
            toaster.show_toast("Downloaded Online Version ",
                                        " ",
                                        icon_path="scrapercon.ico",
                                        duration=5)
            open('LSK Act .pdf', 'wb').write(r.content)
            print(title.string, "rules download complete. ")

            Download1=input ("Do you need the LSK draft rules? [y/n] ").lower()
            if Download1=="yes" or Download1=="y":
                Download1= "https://lsk.or.ke/Downloads/Law%20Society%20of%20Kenya%20(General)%20Regulations%202017-Final%20Draft.pdf"
                r = requests.get(Download1, allow_redirects=True)
                toaster.show_toast("Downloaded Online Version ",
                                        " ",
                                        icon_path="scrapercon.ico",
                                        duration=5)
                open('LSK Draft Rules.pdf', 'wb').write(r.content)
                print(title.string, "rules download complete. ")

    ###
        if actTitle=="Access to information" or actTitle=="information access" or actTitle=="access" or actTitle=="information":
          
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=No.%2031%20of%202016"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
            
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)
            
            ##time = soup.find("select", id="cboPIT")
            ##print("Amendment History")
            ##print(time.text)
            print("\n")
            print(User_Name, ", this is the time stamp of the legislation that matches your keyword: ")
            advo2= "No amendments recorded. Consequential Amendments *few*.  Records Disposal Act-substitute, Public Archives and Documentation Act Service Act: s.5A-Insert, s.6(2,3)-Delete, s.7-Insert, Public Officer Ethics Act: s.41-Insert, Official Secrets Act: s.7-Insert, Statistics Act: s. 11-Insert. "
            list=textwrap3.wrap(advo2, width=90)
            for element in list:
                print(element)
                
            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/A/Access%20to%20Information%20-%20No.%2031%20of%202016/docs/AccesstoInformationActNo31of2016.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("-----Downloaded Online Version-----.")
            print("\n")     
            open('Access to Information .pdf', 'wb').write(r.content)

    ###
        if actTitle=="African re-insurance corporation" or actTitle=="African re-insurance" or actTitle=="re-insurance" or actTitle=="ARC" :
            
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=No.%202%20of%201994"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
           
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)

            ##time = soup.find("select", id="cboPIT")
            ##print("Amendment History")
            ##print(time.text)
            print("\n")
            print(User_Name, ", this is the time stamp of the legislation that matches your keyword: ")
            prov2= "No amendments recorded. "
            list=textwrap3.wrap(prov2, width=85)
            for element in list:
                print(element)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/A/African%20Re-insurance%20Corporation%20Mandatory%20Re-insurance%20Cessions%20Act%20Cap.%20487B/docs/AfricanRe-insuranceCorporationMandatoryRe-insuranceCessionsAct.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("-----Downloaded Online Version-----.")
            open('African re-insurance corporation Act .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Clinical Officers(Training, Registration and Licensing)" or actTitle=="clinical officers" or actTitle=="clinical" or actTitle=="clinic" :
        
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%20260"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
            
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)
            print("\n")
            #time = soup.find("select", id="cboPIT")
            #print("Amendment History")
            #print(time.text)
            print("\n")
            print(User_Name, ", this is the time stamp of the legislation that matches your keyword: ")
            prov2= "No amendments recorded. "
            list=textwrap3.wrap(prov2, width=85)
            for element in list:
                print(element)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/C/Clinical%20Officers%20Training%20Registration%20and%20Licensing%20Act%20Cap.%20260%20-%20No.%209%20of%201988/docs/ClinicalOfficersAct9of1988.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("-----Downloaded Online Version-----.")
            open('Clinical Officers Act .pdf', 'wb').write(r.content)

            print("-"*50)
            clinic3=input ("Do you need the Clinical Officers (Training, Registration and Licensing) subsidiary legislation? [y/n] ").lower()
            if clinic3=="yes" or clinic3=="y":
                clinic3= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/A/Advocates%20Act%20Cap.%2016%20-%20No.%2018%20of%201989/subsidiary%20legislation/docs/AdvocatesAct18of1989_subsidiary.pdf"
                r = requests.get(clinic3, allow_redirects=True)
                clinic4= "Downloading Clinical Officers (Training, Registration and Licensing) subsidiary legislation as *Clinical Officers Regulations*."
                list=textwrap3.wrap(clinic4, width=90)
                for element in list:
                    print(element)
                print("-----Download complete-----")
            print("-"*50)

    ###
        if actTitle=="Coast Development Authority" or actTitle=="coast development" or actTitle=="coast" or actTitle=="coast authority" :
         
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%20449"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
           
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)

            ##time = soup.find("select", id="cboPIT")
            ##print("Amendment History")
            ##print(time.text)
            print("\n")
            print(User_Name, ", this is the time stamp of the legislation that matches your keyword: ")
            prov2= "No amendments recorded. "
            list=textwrap3.wrap(prov2, width=85)
            for element in list:
                print(element)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/C/Coast%20Development%20Authority%20Act%20Cap.%20449%20-%20No.%2020%20of%201990/docs/CoastDevelopmentAuthorityAct20of1990.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("-----Downloaded Online Version-----.")
            open('Coast Development Authority Act .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Climate change" or actTitle=="climate" or actTitle=="carbon emission":
         
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=No.%2011%20of%202016"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
            
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)

            ##time = soup.find("select", id="cboPIT")
            ##print("Amendment History")
            ##print(time.text)
            print("\n")
            print(User_Name, ", this is the time stamp of the legislation that matches your keyword: ")
            prov2= "No amendments recorded. "
            list=textwrap3.wrap(prov2, width=85)
            for element in list:
                print(element)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/C/Climate%20Change%20Act%20-%20No.%2011%20of%202016/docs/ClimateChangeAct11of2016.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("-----Downloaded Online Version-----. ")
            open('Climate Change Act .pdf', 'wb').write(r.content)

    ###
        if actTitle=="College of arms" or actTitle=="college" or actTitle=="arms":
         
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%2098"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
           
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)
            print("\n")
            #time = soup.find("select", id="cboPIT")
            #print("Amendment History")
            #print(time.text)
            print("\n")
            print(User_Name, ", this is the time stamp of the legislation that matches your keyword: ")
            prov2= "Last amendments *few*. November 14 2002 see the Laws of Kenya Rectification Order 2002. s.3(4)-Insert "
            list=textwrap3.wrap(prov2, width=85)
            for element in list:
                print(element)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/C/Climate%20Change%20Act%20-%20No.%2011%20of%202016/docs/ClimateChangeAct11of2016.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("-----Downloaded Online Version-----. ")
            open('College of Arms Act .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Commission for the implementation of the constitution" or actTitle=="cic":
          
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=No.%209%20of%202010"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
            
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)

            ##time = soup.find("select", id="cboPIT")
            ##print("Amendment History")
            ##print(time.text)
            print("\n")
            print(User_Name, ", this is the time stamp of the legislation that matches your keyword: ")
            prov2= "No amendments recorded. "
            list=textwrap3.wrap(prov2, width=85)
            for element in list:
                print(element)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/C/Commission%20for%20the%20Implementation%20of%20the%20Constitution%20Act%20Cap.%205A%20-%20No.%209%20of%202010/docs/CommissionfortheImplementationoftheConstitution9of2010.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("-----Downloaded Online Version-----. ")
            open('Commission for the implementation of the constitution .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Commission on administrative justice" or actTitle=="admin justice" or actTitle=="administrative justice":
         
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=No.%2023%20of%202011"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")

            
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)

            ##time = soup.find("select", id="cboPIT")
            ##print("Amendment History")
            ##print(time.text)
            print("\n")
            print(User_Name, ", this is the time stamp of the legislation that matches your keyword: ")
            prov2= "No amendments recorded. "
            list=textwrap3.wrap(prov2, width=85)
            for element in list:
                print(element)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/C/Commission%20on%20Administrative%20Justice%20Act%20Cap.%20102A%20-%20No.%2023%20of%202011/docs/CommissiononAdministrativeJusticeAct23of2011.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("-----Downloaded Online Version-----. ")
            open('Commission on Administrative Justice .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Commission on Revenue Allocation" or actTitle=="revenue allocation":
         
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=No.%2016%20of%202011"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
         
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)
            print("\n")
            #time = soup.find("select", id="cboPIT")
            #print("Amendment History")
            #print(time.text)
            print("\n")
            print(User_Name, ", this is the time stamp of the legislation that matches your keyword: ")
            prov2= "No amendments recorded. "
            list=textwrap3.wrap(prov2, width=85)
            for element in list:
                print(element)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/C/Commission%20on%20Revenue%20Allocation%20Cap.%205E%20-%20Act%20No.%2016%20of%202011/docs/CommissionOnRevenueAllocationAct16of2011.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("-----Downloaded Online Version-----. ")
            open('Commission on Revenue Allocation .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Commissioners of assize" or actTitle=="assize":
         
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%2012"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
           
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)

            ##time = soup.find("select", id="cboPIT")
            ##print("Amendment History")
            ##print(time.text)
            print("\n")
            print(User_Name, ", this is the time stamp of the legislation that matches your keyword: ")
            prov2= "Last amendments *few*. May 30 1986 see Commissioners of Assize(Amendment) Act 4/1986. s.2-Insert, s.3-Insert, s.5-Repealed, Promissory Oaths Act First Schedule-Insert. "
            list=textwrap3.wrap(prov2, width=85)
            for element in list:
                print(element)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/C/Commissioners%20of%20Assize%20Act%20Cap.%2012%20-%20No.%2038%20of%201954/docs/CommissionersofAssizeAct38of1954.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("-----Downloaded Online Version-----. ")
            open('Commissioners of Assize .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Commissions of Inquiry" or actTitle=="commssions inquiry" or actTitle=="inquiry":
         
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%20102"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
          
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)
            print("\n")
            #time = soup.find("select", id="cboPIT")
            #print("Amendment History")
            #print(time.text)
            print("\n")
            print(User_Name, ", this is the time stamp of the legislation that matches your keyword: ")
            prov2= "Last amendments *few*. August 30 2010 see Commissions of Inquiry (Amendment) Act 5/2010. s.7 substitute.  "
            list=textwrap3.wrap(prov2, width=85)
            for element in list:
                print(element)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/C/Commissions%20of%20Inquiry%20Act%20Cap.%20102%20-%20No.%2011%20of%201962/docs/CommissionsOfInquiryAct11of1962.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("-----Downloaded Online Version-----. ")
            open('Commissions of Inquiry .pdf', 'wb').write(r.content)

    ###
        if actTitle=="community land" or actTitle=="community":
        
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=No.%2027%20of%202016"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
          
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)

            ##time = soup.find("select", id="cboPIT")
            ##print("Amendment History")
            ##print(time.text)
            print("\n")
            print(User_Name, ", this is the time stamp of the legislation that matches your keyword: ")
            prov2= "No amendments recorded.  "
            list=textwrap3.wrap(prov2, width=85)
            for element in list:
                print(element)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/C/Community%20Land%20Act%20-%20No.%2027%20of%202016/docs/CommunityLandAct27of2016.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("-----Downloaded Online Version-----. ")
            open('Community Land .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Community Service Orders" or actTitle=="community service":
          
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=No.%2010%20of%201998"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")

            
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)

            #time = soup.find("select", id="cboPIT")
            #print("Amendment History")
            #print(time.text)
            print("\n")
            print(User_Name, ", this is the time stamp of the legislation that matches your keyword: ")
            prov2= "Recent amendments *few*. May 4 2017 see Statute Law (Miscellaneous Amendments) Act 11/2017. s.2-Insert, s.7(1c)-delete, s.7(1e)-substitute, s.7(1f)-substitute, s.7(1A)-Insert. Consequential repeals *few* see Prisons Act Part 7-repealed, Detention Camps Act-repealed. "
            list=textwrap3.wrap(prov2, width=85)
            for element in list:
                print(element)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/C/Community%20Service%20Orders%20Act%20Cap.%2093%20-%20No.%2010%20of%201998/docs/CommunityServiceOrdersAct10of1998.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("-----Downloaded Online Version-----. ")
            open('Community Service .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Competition" or actTitle=="competition":
        
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=No.%2012%20of%202010"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
         
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)
            print("\n")
            #time = soup.find("select", id="cboPIT")
            #print("Amendment History")
            #print(time.text)
            print("\n")
            print(User_Name, ", this is the time stamp of the legislation that matches your keyword: ")
            prov2= "Recent amendments *few*. May 4 2017 see schedule to Statute Law (Miscellaneous Amendments) Act 11/2017.  Para. 3(4)-substitute. "
            list=textwrap3.wrap(prov2, width=85)
            for element in list:
                print(element)

            Download= "http://www.kenyalaw.org/lex//actview.xql?actid=No.%2012%20of%202010"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("-----Downloaded Online Version-----. ")
            open('Competition Act .pdf', 'wb').write(r.content)
            print("-"*50)
            compe3=input ("Do you need the Competition subsidiary legislation? [y/n] ").lower()
            if compe3=="yes" or compe3=="y":
                compe4= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/C/Competition%20Act%20Cap.%20504%20-%20No.%2012%20of%202010/docs/CompetitionAct12of2010.pdf"
                r = requests.get(compe4, allow_redirects=True)
                clinic4= "Downloading Competition subsidiary legislation as *Competition Regulations*."
                list=textwrap3.wrap(clinic4, width=90)
                for element in list:
                    print(element)
                open('Competition Regulations .pdf', 'wb').write(r.content)
                print("-----Download complete-----")
            print("-"*50)
            print("\n")
            compe3=input ("Would you like to view the latest decisions by the Competition Authority of Kenya? [y/n] ").lower()
            if compe3=="yes" or compe3=="y":
                compe4= "https://www.cak.go.ke/images/new/Key_CAK_Decisions_June_12_2018.pdf"
                r = requests.get(compe4, allow_redirects=True)
                print("\n")
                compe8= "Downloading latest decisions by the Competition Authority of Kenya as *CAK Decisions as at June 12 2018*."
                list=textwrap3.wrap(compe8, width=90)
                for element in list:
                    print(element)
                print("-----Download complete-----")
            open('CAK Decisions as at June 12 2018.pdf', 'wb').write(r.content)
            print("-"*50)
            print("\n")
            compe5=input ("How about the 2017 CAK decisions? [y/n] ").lower()
            if compe5=="yes" or compe5=="y":
                compe6= "https://www.cak.go.ke/images/new/CAK_DECISIONS_29_AUG_2017.pdf"
                r = requests.get(compe6, allow_redirects=True)
                print("\n")
                compe7= "Downloading August 2017 decisions by the Competition Authority of Kenya as *CAK Decisions as at August 29 2017*."
                list=textwrap3.wrap(compe7, width=90)
                for element in list:
                    print(element)
                open('CAK Decisions as at August 29 2017.pdf', 'wb').write(r.content)
                print("-----Download complete-----")
            print("-"*50)

    ###
        if actTitle=="Compounding of Potable Spirits" or actTitle=="compounding of portable spirits" or actTitle=="compounding portable spirits" or actTitle=="portable spirits" or actTitle=="Compounding of Portable Spirits" or actTitle=="potable spirits" or actTitle=="spirits":
         
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%20123"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
           
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)

            ##time = soup.find("select", id="cboPIT")
            ##print("Amendment History")
            ##print(time.text)
            print("\n")
            print(User_Name, ", this is the time stamp of the legislation that matches your keyword: ")
            prov2= "Last amendments *few*. July 12 1966 see Statute Law (Miscellaneous Amendments) Act 21/1966. s.3(v)-Substitute. "
            list=textwrap3.wrap(prov2, width=85)
            for element in list:
                print(element)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/C/Compounding%20of%20Portable%20Spirits%20Act%20Cap.%20123%20-%20No.%2016%20of%201961/docs/CompoundingofPotableSpiritsAct16of1961.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("-----Downloaded Online Version-----. ")
            open('Compounding of Potable Spirits .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Consolidated Bank of Kenya" or actTitle=="consolidated bank" or actTitle=="consolidated" or actTitle=="consolidate" or actTitle=="Consolidated bank of kenya":
         
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=No.%205%20of%201991"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
           
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)

            ##time = soup.find("select", id="cboPIT")
            ##print("Amendment History")
            ##print(time.text)
            print("\n")
            print(User_Name, ", this is the time stamp of the legislation that matches your keyword: ")
            prov2= "No amendments recorded. "
            list=textwrap3.wrap(prov2, width=85)
            for element in list:
                print(element)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/C/Consolidated%20Bank%20of%20Kenya%20Act%20%20Cap.%20488C%20-%20No.%205%20of%201991/docs/ConsolidatedBankofKenyaAct5of1991.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("-----Downloaded Online Version-----. ")
            open('Cosolidated Bank .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Constituencies Development Fund" or actTitle=="cdf":
         
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=No.%2030%20of%202013"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
           
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)

            ##time = soup.find("select", id="cboPIT")
            ##print("Amendment History")
            ##print(time.text)
            print("\n")
            print(User_Name, ", this is the time stamp of the legislation that matches your keyword: ")
            prov2= "No amendments recorded.  Consequential amendments *few*. Constituencies Development Fund Act 2003-Repealed"
            list=textwrap3.wrap(prov2, width=85)
            for element in list:
                print(element)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/C/Constituencies%20Development%20Fund%20Act%20No.%2030%20of%202013/docs/Constituencies%20Development%20Fund%20Act%20No.%2030%20of%202013.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("-----Downloaded Online Version-----. ")
            open('CDF .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Consumer Protection" or actTitle=="consumer" or actTitle=="Article 46" or actTitle=="consumer protection":
         
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=No.%2046%20of%202012"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
            
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)

            ##time = soup.find("select", id="cboPIT")
            ##print("Amendment History")
            ##print(time.text)
            print("\n")
            print(User_Name, ", this is the time stamp of the legislation that matches your keyword: ")
            prov2= "Recent amendments *few*. September 20 2016 see Finance Act 38/2016. s.62-Insert."
            list=textwrap3.wrap(prov2, width=85)
            for element in list:
                print(element)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/C/Constituencies%20Development%20Fund%20Act%20No.%2030%20of%202013/docs/Constituencies%20Development%20Fund%20Act%20No.%2030%20of%202013.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("-----Downloaded Online Version-----. ")
            open('Consumer Protection .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Contracts Restraint of Trade" or actTitle=="contracts restraint of trade" or actTitle=="restraint of trade":
        
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%2024"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
           
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)

            ##time = soup.find("select", id="cboPIT")
            ##print("Amendment History")
            ##print(time.text)
            print("\n")
            print(User_Name, ", this is the time stamp of the legislation that matches your keyword: ")
            prov2= "No amendments recorded."
            list=textwrap3.wrap(prov2, width=85)
            for element in list:
                print(element)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/C/Constituencies%20Development%20Fund%20Act%20No.%2030%20of%202013/docs/Constituencies%20Development%20Fund%20Act%20No.%2030%20of%202013.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("I downloaded a copy just in case.")
            open('Contracts in Restraint of Trade .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Controller of Budget" or actTitle=="controller of budget" or actTitle=="budget controller" or actTitle=="budget":
        
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=No.%2026%20of%202016"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
            
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)

            ##time = soup.find("select", id="cboPIT")
            ##print("Amendment History")
            ##print(time.text)
            print("\n")
            print(User_Name, ", this is the time stamp of the legislation that matches your keyword: ")
            prov2= "No amendments recorded. Consequential amendments *few*. Independent Offices (Appointment) Act-repealed. "
            list=textwrap3.wrap(prov2, width=85)
            for element in list:
                print(element)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/C/Controller%20of%20Budget%20-%20No.%2026%20of%202016/docs/ControllerofBudgetAct26of2016.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("-----Downloaded Online Version-----. ")
            open('Controller of Budget .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Customs and Excise" or actTitle=="customs and excise" or actTitle=="customs" or actTitle=="excise" or actTitle=="custom excise" or actTitle=="customs excise":
          
            print("Let's be patient. This is a voluminous one. ")
            
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%20472"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
           
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)
            print("Accurate results take time, just like this loader")

            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)

            ##time = soup.find("select", id="cboPIT")
            ##print("Amendment History")
            ##print(time.text)
            print("\n")
            print(User_Name, ", this is the time stamp of the legislation that matches your keyword: ")
            prov2= "Last amendments *many*. See Finance Act 38/2013. June 18 2013, s.2-Insert. June 18 2013, s.91-Substitute. January 1 2014, s.91A-Insert. January 1 2014, s.92-Insert. July 1 2013, s.117A-Insert. January 1 2014, Fifth Schedule, Part III, Item 7-substitute, para. 9-Insert. February 1 2013, s.137(1)-delete, 137(3,4,5)-Insert. "
            list=textwrap3.wrap(prov2, width=85)
            for element in list:
                print(element)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/C/Customs%20and%20Excise%20Act%20Cap.%20472%20-%20No.%2010%20of%201978/docs/CustomsandExciseAct10of1978.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("-----Downloaded Online Version-----. ")
            open('Customs and Excise .pdf', 'wb').write(r.content)

            Download1=input ("Do you need the Customs and Excise Subsidiary Legislation? [y/n] ").lower()
            if Download1=="yes" or Download1=="y":
                Download1= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/C/Customs%20and%20Excise%20Act%20Cap.%20472%20-%20No.%2010%20of%201978/subsidiary%20legislation/docs/CustomsandExciseAct10of1978_subsidiary.pdf"
                r = requests.get(Download1, allow_redirects=True)
                Download1= "Downloading Customs and Excise Subsidiary Legislation *Customs Regulations*."
                list=textwrap3.wrap(Download1, width=90)
                for element in list:
                    print(element)
                print("----Download complete----.")
                open('Customs Regulations.pdf', 'wb').write(r.content)

            Download1=input ("Are you interested in the East African Community Custom Management Act? [y/n] ").lower()
            if Download1=="yes" or Download1=="y":
                Download1= "http://www.kra.go.ke/notices/pdf2018/East%20African%20Community%20Customs%20Management%20Act%20(revised%2030th%20June%202017).pdf"
                r = requests.get(Download1, allow_redirects=True)
                Download1= "Downloading EAC Customs Management Act as *EAC Customs Act*."
                list=textwrap3.wrap(Download1, width=90)
                for element in list:
                    print(element)
                print("----Download complete----.")
                open('EAC Customs Act.pdf', 'wb').write(r.content)

    ###
        if actTitle=="SUPPLIES PRACTITIONERS MANAGEMENT" or actTitle=="supplies practitioners management" or actTitle=="supplies practitioners" or actTitle=="supplies":
          
            print("Let's be patient. This is a voluminous one. ")
            
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=No.%2017%20of%202007"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")

            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)
            print("\n")
            #time = soup.find("select", id="cboPIT")
            #print("Amendment History")
            #print(time.text)
            print("\n")
            print(User_Name, ", this is the time stamp of the legislation that matches your keyword: ")
            prov2= "Last amendments *many*. July 23 2009 see Statute Law (Miscellaneous Amendment) Act 6/2009. See also nexus with Public Procurement and Asset Disposal Act. "
            list=textwrap3.wrap(prov2, width=85)
            for element in list:
                print(element)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/S/Supplies%20Practitioners%20Management%20Act%20No.%2017%20of%202007/docs/SuppliesPractitionersManagementAct17of2007.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("-----Downloaded Online Version-----. ")
            open('Supplies Practitioners Management.pdf', 'wb').write(r.content)
            print("\n")
            Download1=input ("Do you need the Supplies Practitioners Management Subsidiary Legislation? [y/n] ").lower()
            if Download1=="yes" or Download1=="y":
                Download1= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/S/Supplies%20Practitioners%20Management%20Act%20No.%2017%20of%202007/subsidiary%20legislation/docs/SuppliesPractitionersManagementAct17of2007_subsidiary.pdf"
                r = requests.get(Download1, allow_redirects=True)
                Download1= "Downloading Supplies Practitioners Management Subsidiary Legislation *Procurement Officers Regulations*."
                list=textwrap3.wrap(Download1, width=90)
                for element in list:
                    print(element)
                print("----Download complete----.")
                open('Procurement Officers.pdf', 'wb').write(r.content)
            print("\n")
            Download1=input ("Are you interested in Recent Press regarding the subject? [y/n] ").lower()
            if Download1=="yes" or Download1=="y":
                Download1= "http://media.reelforge.com/player/view/?file=/2018/06/06/DNT_20180606_5VIJK5LC5CQ.pdf&name=download_Professional_bodies_support_Uhuru_order,_procurement_officers_comply_Daily_Nation__06-Jun-2018_:_Page_5"
                r = requests.get(Download1, allow_redirects=True)
                Download1= "Downloading as *Recent News*."
                list=textwrap3.wrap(Download1, width=90)
                for element in list:
                    print(element)
                print("----Download complete----.")
                open('Recent News.pdf', 'wb').write(r.content)
    ###
        if actTitle=="Physical Planning" or actTitle=="physical planning" or actTitle=="Physical planning" or actTitle=="PHYSICAL PLANNING":
            
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=No.%206%20of%201996"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
           
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)
            #time = soup.find("select", id="cboPIT")
            #print("Amendment History")
            #print(time.text)
            print("\n")
            print(User_Name, ", this is the time stamp of the legislation that matches your keyword: ")
            prov2= "No amendments recorded. "
            list=textwrap3.wrap(prov2, width=85)
            for element in list:
                print(element)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/P/Physical%20Planning%20Act%20Cap.%20286%20-%20No.%206%20of%201996/docs/PhysicalPlanningAct6of1996.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("-----Downloaded Online Version-----.")
            print("\n")
            open('Physical Planning .pdf', 'wb').write(r.content)

    ###
        if actTitle=="Survey" or actTitle=="survey" or actTitle=="SURVEY" :
            
            actTitle1= ["http://www.kenyalaw.org/lex//actview.xql?actid=CAP.%20299"]

            for urls in actTitle1:
                response = requests.get(urls)
                soup = BeautifulSoup(response.content, "lxml")
           
            title= soup.find("div", attrs={"class":"act-title"})
            for title in title:
                print(title.string)

            commencement=soup.find_all("div", class_="act-commencement")
            for commencement in commencement:
                print(commencement.text)
            #time = soup.find("select", id="cboPIT")
            #print("Amendment History")
            #print(time.text)
            print("\n")
            print(User_Name, "this is the time stamp of the legislation that matches your keyword:")
            print("\n")
            prov2= "No amendments recorded. "
            list=textwrap3.wrap(prov2, width=85)
            for element in list:
                print(element)

            Download= "http://www.kenyalaw.org/lex/rest//db/kenyalex/Kenya/Legislation/English/Acts%20and%20Regulations/S/Survey%20Act%20Cap.%20299%20-%20Act%20No.%2025%20of%201961/docs/SurveyAct25of1961.pdf"
            r = requests.get(Download, allow_redirects=True)
            print("\n")
            print("-----Downloaded Online Version-----.")
            print("\n")
            open('Survey.pdf', 'wb').write(r.content)

    ###
        if actTitle=="data protection" or actTitle=="DATA PROTECTION" or actTitle=="DATA" or actTitle=="data":
            dpa1="http://kenyalaw.org/kl/fileadmin/pdfdownloads/bills/2018/DataProtectionBill_2018.pdf"
            r = requests.get(dpa1, allow_redirects=True)
            print("\n")
            print("Downloading *Data Protection Bill, 2018* ")
            print("----Download Complete----.")
            open('Data Protection Bill .pdf', 'wb').write(r.content)
            print("\n")
            print("-"*50)
            dpa2="https://www.gov.uk/data-protection/print"
            r = requests.get(dpa2, allow_redirects=True)
            print ("Downloading *UK GDPR guide* as UK equivalent. ")
            print("----Download Complete----.")
            open('UK GDPR guide .pdf', 'wb').write(r.content)
            print("\n")
            print("-"*50)
        
    #Ask user for a different action
        if actTitle in actTitle:
            print("-"*50)
            restart=input ("Pick a different search activity? [y/n] ").lower()
            if restart=="yes" or restart=="y":
                print("\n")
                main()
            else:
                from datetime import datetime, date
                now= datetime.now().time()
                if date.today().weekday()== 0:
                    if now.hour == 1 or now.hour == 2 or now.hour == 3 or now.hour == 4 or now.hour == 5 or now.hour == 6 or now.hour == 7 or now.hour == 8 or now.hour == 9:
                        print("Enjoy the rest of your Monday morning", User_Name, "------")
                    if now.hour==10 or now.hour == 11 or now.hour == 12:
                        print("Have a good day", User_Name, )
                    if now.hour==13 or now.hour == 14 or now.hour == 15 or now.hour == 16:
                        print("Have good afternoon", User_Name)
                if date.today().weekday()== 1:
                    if now.hour == 1 or now.hour == 2 or now.hour == 3 or now.hour == 4 or now.hour == 5 or now.hour == 6 or now.hour == 7 or now.hour == 8 or now.hour == 9:
                        print("Enjoy the rest of your Monday morning", User_Name, "------")
                    if now.hour==10 or now.hour == 11 or now.hour == 12:
                        print("Have a good day", User_Name, )
                    if now.hour==13 or now.hour == 14 or now.hour == 15 or now.hour == 16:
                        print("Have good afternoon", User_Name)
                if date.today().weekday()== 2:
                    if now.hour == 1 or now.hour == 2 or now.hour == 3 or now.hour == 4 or now.hour == 5 or now.hour == 6 or now.hour == 7 or now.hour == 8 or now.hour == 9:
                        print("Enjoy the rest of your Monday morning", User_Name, "------")
                    if now.hour==10 or now.hour == 11 or now.hour == 12:
                        print("Have a good day", User_Name, )
                    if now.hour==13 or now.hour == 14 or now.hour == 15 or now.hour == 16:
                        print("Have good afternoon", User_Name)
                if date.today().weekday()== 3:
                    if now.hour == 1 or now.hour == 2 or now.hour == 3 or now.hour == 4 or now.hour == 5 or now.hour == 6 or now.hour == 7 or now.hour == 8 or now.hour == 9:
                        print("Enjoy the rest of your Monday morning", User_Name, "------")
                    if now.hour==10 or now.hour == 11 or now.hour == 12:
                        print("Have a good day", User_Name, )
                    if now.hour==13 or now.hour == 14 or now.hour == 15 or now.hour == 16:
                        print("Have good afternoon", User_Name)
                if date.today().weekday()== 4:
                    if now.hour == 1 or now.hour == 2 or now.hour == 3 or now.hour == 4 or now.hour == 5 or now.hour == 6 or now.hour == 7 or now.hour == 8 or now.hour == 9:
                        print("Enjoy the rest of your morning", User_Name, "------")
                    if now.hour==10 or now.hour == 11 or now.hour == 12:
                        print("Have a good day", User_Name, )
                    if now.hour==13 or now.hour == 14 or now.hour == 15 or now.hour == 16:
                        print("Have good afternoon", User_Name)
                    if now.hour==17 or now.hour == 18 or now.hour == 19 or now.hour == 20 or now.hour == 21 or now.hour == 22 or now.hour == 23:
                        print("Have a great weekend", User_Name)
                if date.today().weekday()== 5:
                    if now.hour == 1 or now.hour == 2 or now.hour == 3 or now.hour == 4 or now.hour == 5 or now.hour == 6 or now.hour == 7 or now.hour == 8 or now.hour == 9:
                        print("Enjoy the rest of your morning", User_Name, "------")
                    if now.hour==10 or now.hour == 11 or now.hour == 12:
                        print("Have a good day", User_Name, )
                    if now.hour==13 or now.hour == 14 or now.hour == 15 or now.hour == 16:
                        print("Have good afternoon", User_Name)
                    if now.hour==17 or now.hour == 18 or now.hour == 19 or now.hour == 20 or now.hour == 21 or now.hour == 22 or now.hour == 23:
                        print("Enjoy your weekend", User_Name)
                if date.today().weekday()== 6:
                    if now.hour == 1 or now.hour == 2 or now.hour == 3 or now.hour == 4 or now.hour == 5 or now.hour == 6 or now.hour == 7 or now.hour == 8 or now.hour == 9:
                        print("Enjoy the rest of your morning", User_Name, "------")
                    if now.hour==10 or now.hour == 11 or now.hour == 12:
                        print("Have a good day", User_Name, )
                    if now.hour==13 or now.hour == 14 or now.hour == 15 or now.hour == 16:
                        print("Have good afternoon", User_Name)
                    if now.hour==17 or now.hour == 18 or now.hour == 19 or now.hour == 20 or now.hour == 21 or now.hour == 22 or now.hour == 23:
                        print("Have a great week ahead", User_Name)
                print("\n")
                print ("Activity Summary: ")
                print("> Researcher: ", User_Name)
                print("> Legislation keyword: ", "'", actTitle, "'")
                print("\n")
                import csv
                with open("Legislation_Search_Activity.csv", "a", encoding="utf-8") as csv_file:
                    writer = csv.writer(csv_file)
                    writer.writerow(["Date", "Researcher", "Legislation", "Summary"])
                    Date=date.today()
                    Researcher=User_Name
                    Legislation=actTitle
                    Summary=Amending_Acts.text
                    writer.writerow([Date, Researcher, Legislation, Summary])
     
#DEPENDING ON USER INPUT, CODE WILL EITHER END OR RESTART HERE
print("\n")
main()
input("Press 'Enter' to exit ")
sys.exit()


###---THE END---###




    
        




    




